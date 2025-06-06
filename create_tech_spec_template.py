import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_tech_spec_hydraulic_press_template():
    content = """
SPÉCIFICATIONS TECHNIQUES – PRESSE HYDRAULIQUE

Modèle de la Presse : [Modèle de la Presse]
Date de Spécification : [Date]

## 1. Caractéristiques Générales
- Force Nominale (Tonnes) : [Force Nominale]
- Course du Coulisseau (mm) : [Course du Coulisseau]
- Ouverture Maximale (mm) : [Ouverture Maximale] (Distance entre table et coulisseau)
- Dimensions de la Table (mm x mm) : [Largeur Table] x [Profondeur Table]
- Dimensions du Coulisseau (mm x mm) : [Largeur Coulisseau] x [Profondeur Coulisseau]
- Vitesse d'Approche (mm/s) : [Vitesse d'Approche]
- Vitesse de Travail (mm/s) : [Vitesse de Travail]
- Vitesse de Retour (mm/s) : [Vitesse de Retour]
- Hauteur de Travail (mm) : [Hauteur de Travail] (Distance du sol à la table)
- Dimensions Générales (L x l x H mm) : [Longueur] x [Largeur] x [Hauteur]
- Poids Approximatif (kg) : [Poids Approximatif]

## 2. Structure
- Type de Bâti : [Type de Bâti : Col de cygne, 2 montants, 4 colonnes, etc.]
- Matériau du Bâti : [Matériau du Bâti, e.g., Acier S355JR mécano-soudé]
- Table : [Description de la table : fixe, mobile, avec rainures en T (dimensions, nombre)]
- Coulisseau : [Description du coulisseau : guidage (prismatique, cylindrique), avec rainures en T]

## 3. Groupe Hydraulique
- Type de Pompe : [Type de Pompe : à pistons, à engrenages, etc.]
- Puissance Moteur Pompe (kW/CV) : [Puissance Moteur Pompe]
- Pression Maximale de Service (bar) : [Pression Maximale]
- Capacité du Réservoir d'Huile (litres) : [Capacité Réservoir]
- Type d'Huile Recommandée : [Type d'Huile]
- Filtration : [Description du système de filtration]
- Refroidissement : [Type de système de refroidissement : air/huile, eau/huile, aucun]
- Composants : [Marque des principaux composants : distributeurs, vannes, etc. e.g., Bosch Rexroth, Parker]

## 4. Commandes Électriques
- Type d'Automate (PLC) : [Marque et modèle de l'automate, e.g., Siemens S7-1200]
- Pupitre de Commande : [Description : écran tactile (taille), boutons poussoirs, bimanuelle]
- Tension d'Alimentation : [Tension V] / [Fréquence Hz] / [Phases Ph]
- Puissance Installée Totale (kW) : [Puissance Totale]
- Armoire Électrique : [Description : position, indice de protection IPXX]

## 5. Sécurité
- Dispositifs de Sécurité : [Barrières immatérielles (type, catégorie), portes de protection (type, verrouillage), cales de sécurité manuelles/hydrauliques, etc.]
- Conformité Normes : [Normes respectées, e.g., CE EN ISO 16092-1, EN 60204-1]

## 6. Options Incluses
- [Option 1 : Description]
- [Option 2 : Description]
- [Option 3 : Description]

## 7. Documentation Fournie
- Manuel d'Utilisation et de Maintenance (Langue : [Langue Documentation])
- Schémas Électriques et Hydrauliques
- Déclaration de Conformité CE

## 8. Conditions Commerciales
- Garantie : [Durée et conditions de garantie]
- Conditions de Paiement : [Conditions de paiement]
- Délai de Livraison : [Délai de livraison estimé]

SIGNATURES
"""

    doc = Document()

    # Set default font for the document (optional)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = doc.add_heading("SPÉCIFICATIONS TECHNIQUES – PRESSE HYDRAULIQUE", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add content
    lines = content.strip().split('\n')
    is_first_content_line = True # To skip initial empty lines from strip() if any before title

    for line in lines:
        stripped_line = line.strip()
        if not stripped_line and not is_first_content_line: # Add an empty paragraph for spacing
            doc.add_paragraph("")
            continue
        elif not stripped_line and is_first_content_line:
            continue # Skip leading empty lines after title

        is_first_content_line = False

        if stripped_line.startswith("## "):
            heading_text = stripped_line[3:] # Remove "## "
            doc.add_heading(heading_text, level=2)
        elif stripped_line == "SIGNATURES":
             p = doc.add_paragraph()
             p.alignment = WD_ALIGN_PARAGRAPH.CENTER
             run = p.add_run(stripped_line)
             run.bold = True
             run.font.size = Pt(12)
        else:
            # For list items starting with '-'
            if stripped_line.startswith("- "):
                p = doc.add_paragraph(style='ListBullet')
                p.text = stripped_line[2:] # Remove "- "
                # Set indent for bullet points if needed (more advanced)
                # p.paragraph_format.left_indent = Inches(0.25)
                # p.paragraph_format.first_line_indent = Inches(-0.25)
            else:
                 doc.add_paragraph(stripped_line)


    # Ensure directory exists
    output_dir = "templates/fr/"
    os.makedirs(output_dir, exist_ok=True)

    # The problem description asks for specification_technique_template.docx
    # but the ls output showed specification_technique_template.xlsx
    # I will use the .docx extension as requested.
    file_path = os.path.join(output_dir, "specification_technique_template.docx")
    doc.save(file_path)
    print(f"Document saved to {file_path}")

if __name__ == "__main__":
    create_tech_spec_hydraulic_press_template()
