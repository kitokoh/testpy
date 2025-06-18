# -*- coding: utf-8 -*-
import os
import json
import sys
import logging # Added logging
from datetime import datetime

from PyQt5.QtCore import QStandardPaths, QCoreApplication, QUrl
from PyQt5.QtGui import QDesktopServices # Added for generate_pdf_for_document
from PyQt5.QtWidgets import QMessageBox

from docx import Document
import db as db_manager # For generate_pdf_for_document
from html_editor import HtmlEditor # For generate_pdf_for_document
from html_to_pdf_util import convert_html_to_pdf # For generate_pdf_for_document

# MAIN_APP_ROOT_DIR import removed as app_root_dir is now passed to generate_pdf_for_document

# --- Configuration Constants ---
CONFIG_DIR_NAME = "ClientDocumentManager"
CONFIG_FILE_NAME = "config.json"
# TEMPLATES_SUBDIR = "templates" # Removed as no longer used in utils.py
# CLIENTS_SUBDIR = "clients"   # Removed as no longer used in utils.py

# --- Configuration Functions ---
def get_config_dir():
    config_dir_path = os.path.join(
        QStandardPaths.writableLocation(QStandardPaths.AppConfigLocation),
        CONFIG_DIR_NAME
    )
    os.makedirs(config_dir_path, exist_ok=True)
    return config_dir_path

def get_config_file_path():
    return os.path.join(get_config_dir(), CONFIG_FILE_NAME)

def load_config(app_root_dir, default_templates_dir, default_clients_dir):
    """
    Loads configuration from JSON file.
    Requires app_root_dir to construct default paths if config file doesn't exist.
    """
    config_path = get_config_file_path()
    if os.path.exists(config_path):
        try:
            with open(config_path, "r", encoding="utf-8") as f:
                return json.load(f)
        except (IOError, json.JSONDecodeError) as e:
            print(f"Error loading config: {e}. Using defaults.")

    # If config file doesn't exist or is invalid, return defaults
    return {
        "templates_dir": default_templates_dir, # Use passed default
        "clients_dir": default_clients_dir,     # Use passed default
        "database_path": os.path.join(app_root_dir, "app_data.db"),
        "language": "fr",
        "smtp_server": "",
        "smtp_port": 587,
        "smtp_user": "",
        "smtp_password": "",
        "default_reminder_days": 30
    }

def save_config(config_data):
    try:
        config_path = get_config_file_path()
        with open(config_path, "w", encoding="utf-8") as f:
            json.dump(config_data, f, indent=4, ensure_ascii=False)
    except IOError as e:
        QMessageBox.warning(None, QCoreApplication.translate("utils.save_config", "Erreur de Configuration"),
                            QCoreApplication.translate("utils.save_config", "Impossible d'enregistrer la configuration: {0}").format(e))

def is_first_launch(app_root_dir, default_templates_dir, default_clients_dir):
    """Checks if this is the first launch of the application."""
    config = load_config(app_root_dir, default_templates_dir, default_clients_dir)
    return not config.get("initial_setup_complete", False)

def mark_initial_setup_complete(app_root_dir, default_templates_dir, default_clients_dir):
    """Marks that the initial setup has been completed."""
    config = load_config(app_root_dir, default_templates_dir, default_clients_dir)
    config["initial_setup_complete"] = True
    save_config(config)

# --- DOCX Population Logic ---
def populate_docx_template(docx_path, client_data):
    """
    Populates a .docx template with client data using placeholders.
    Placeholders should be in the format {{PLACEHOLDER_NAME}}.
    """
    if not os.path.exists(docx_path):
        logging.error(f"DOCX template file not found: {docx_path}")
        QMessageBox.critical(None, "Template Error", f"DOCX template file not found: {docx_path}")
        return # Implicitly returns None

    try:
        document = Document(docx_path)
        placeholders = {
            "{{CLIENT_NAME}}": client_data.get('client_name', ''),
            "{{PROJECT_ID}}": client_data.get('project_identifier', ''),
            "{{COMPANY_NAME}}": client_data.get('company_name', ''),
            "{{NEED}}": client_data.get('need', ''), # This was 'primary_need_description' in EditClientDialog, ensure consistency
            "{{COUNTRY}}": client_data.get('country', ''),
            "{{CITY}}": client_data.get('city', ''),
            "{{PRICE}}": str(client_data.get('price', 0)),
            "{{DATE}}": datetime.now().strftime('%Y-%m-%d'),
            "{{STATUS}}": client_data.get('status', ''),
            "{{SELECTED_LANGUAGES}}": ", ".join(client_data.get('selected_languages', []) if isinstance(client_data.get('selected_languages'), list) else str(client_data.get('selected_languages', '')).split(',')),
            "{{NOTES}}": client_data.get('notes', ''),
            "{{CREATION_DATE}}": client_data.get('creation_date', ''),
            "{{CATEGORY}}": client_data.get('category', ''),
            "{{PRIMARY_CONTACT_NAME}}": "", # Needs logic to fetch primary contact
        }

        for para in document.paragraphs:
            for key, value in placeholders.items():
                if key in para.text:
                    new_text = para.text.replace(key, str(value)) # Ensure value is string
                    if para.text != new_text:
                         para.text = new_text

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for key, value in placeholders.items():
                            if key in para.text:
                                new_text = para.text.replace(key, str(value)) # Ensure value is string
                                if para.text != new_text:
                                    para.text = new_text
        document.save(docx_path)
        logging.info(f"DOCX template populated: {docx_path}") # Changed to logging.info
    except Exception as e:
        logging.error(f"Error populating DOCX template {docx_path}: {e}", exc_info=True) # Added exc_info
        QMessageBox.critical(None, "Population Error", f"Error populating DOCX template {docx_path}: {e}")
        # Optionally, re-raise if the caller needs to handle it, or just return to signify failure.
        # For now, matching the new pattern of informing user and returning.
        return # Implicitly returns None

# --- PDF Generation Logic ---
def generate_pdf_for_document(source_file_path: str, client_info: dict, app_root_dir:str, parent_widget=None, target_language_code: str = "fr") -> str | None:
    """
    Generates a PDF for a given source document (HTML, XLSX, DOCX).
    Uses app_root_dir for resolving paths to resources like default logos if needed.
    Includes logic to check for product language match for proforma invoices.
    """
    if not os.path.exists(source_file_path):
        logging.error(f"Source HTML file not found for PDF generation: {source_file_path}")
        QMessageBox.critical(parent_widget, "File Error", f"Source HTML file not found for PDF generation: {source_file_path}")
        return None

    if not client_info or 'client_id' not in client_info:
        QMessageBox.warning(parent_widget, QCoreApplication.translate("utils.generate_pdf", "Erreur Client"),
                            QCoreApplication.translate("utils.generate_pdf", "ID Client manquant. Impossible de générer le PDF."))
        return None

    file_name, file_ext = os.path.splitext(os.path.basename(source_file_path))
    current_date_str = datetime.now().strftime("%Y%m%d")
    output_pdf_filename = f"{file_name}_{current_date_str}.pdf"
    output_pdf_path = os.path.join(os.path.dirname(source_file_path), output_pdf_filename)

    # Fetch default company ID for context data
    default_company_obj = db_manager.get_default_company()
    default_company_id = default_company_obj['company_id'] if default_company_obj else None
    if not default_company_id and file_ext.lower() == '.html': # Only warn if it's an HTML doc where seller info is crucial
         QMessageBox.information(parent_widget, QCoreApplication.translate("utils.generate_pdf", "Avertissement"),
                                QCoreApplication.translate("utils.generate_pdf", "Aucune société par défaut n'est définie. Les détails du vendeur peuvent être manquants."))

    # Get document context data, now including target_language_code
    # Assuming client_info contains 'client_id'
    # The 'company_id' for get_document_context_data is the seller/our company.
    document_context = db_manager.get_document_context_data(
        client_id=client_info['client_id'],
        company_id=default_company_id, # This should be the ID of "our" company
        target_language_code=target_language_code,
        project_id=client_info.get('project_id'), # Pass project_id if available in client_info
        additional_context={'document_title': file_name} # Example additional context
    )

    if file_ext.lower() == '.html':
        # "No Products" Check for proforma invoices
        # A more robust check might involve looking at a template type if available
        if "proforma" in file_name.lower() or "invoice" in file_name.lower(): # Heuristic check
            products_in_target_lang = [p for p in document_context.get('products', []) if p.get('is_language_match')]
            if not document_context.get('products'): # No products linked at all
                 QMessageBox.information(parent_widget, QCoreApplication.translate("utils.generate_pdf", "Information"),
                                        QCoreApplication.translate("utils.generate_pdf", f"Aucun produit n'est lié à ce client/projet. La génération du PDF pour '{file_name}' est annulée."))
                 return None
            elif not products_in_target_lang:
                lang_name = target_language_code # Ideally, get full language name
                QMessageBox.information(parent_widget, QCoreApplication.translate("utils.generate_pdf", "Information"),
                                        QCoreApplication.translate("utils.generate_pdf", f"Aucun produit trouvé en langue '{lang_name}' pour ce client/projet. La génération du PDF pour '{file_name}' est annulée."))
                return None

        try:
            with open(source_file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()

            # HtmlEditor.populate_html_content now uses the enhanced context
            processed_html = HtmlEditor.populate_html_content(html_content, document_context) # Pass the full context

            base_url = QUrl.fromLocalFile(os.path.dirname(source_file_path)).toString()
            pdf_bytes = convert_html_to_pdf(processed_html, base_url=base_url)

            if pdf_bytes:
                with open(output_pdf_path, 'wb') as f_pdf:
                    f_pdf.write(pdf_bytes)
                QMessageBox.information(parent_widget, QCoreApplication.translate("utils.generate_pdf", "Succès PDF"),
                                        QCoreApplication.translate("utils.generate_pdf", "PDF généré avec succès:\n{0}").format(output_pdf_path))
                return output_pdf_path
            else:
                QMessageBox.warning(parent_widget, QCoreApplication.translate("utils.generate_pdf", "Erreur PDF"),
                                    QCoreApplication.translate("utils.generate_pdf", "La conversion HTML en PDF a échoué. Le contenu PDF résultant était vide."))
                return None
        except Exception as e:
            QMessageBox.critical(parent_widget, QCoreApplication.translate("utils.generate_pdf", "Erreur HTML vers PDF"),
                                 QCoreApplication.translate("utils.generate_pdf", "Erreur lors de la génération du PDF à partir du HTML:\n{0}").format(str(e)))
            return None
    elif file_ext.lower() in ['.xlsx', '.docx']:
        QMessageBox.information(parent_widget, QCoreApplication.translate("utils.generate_pdf", "Fonctionnalité non disponible"),
                                QCoreApplication.translate("utils.generate_pdf", "La génération PDF directe pour les fichiers {0} n'est pas supportée.\nVeuillez utiliser la fonction 'Enregistrer sous PDF' ou 'Exporter vers PDF' de l'application correspondante.").format(file_ext.upper()))
        return None
    else:
        QMessageBox.warning(parent_widget, QCoreApplication.translate("utils.generate_pdf", "Type de fichier non supporté"),
                            QCoreApplication.translate("utils.generate_pdf", "La génération PDF n'est pas supportée pour les fichiers de type '{0}'.").format(file_ext))
        return None

# Note: APP_ROOT_DIR is problematic here.
# If utils.py is at the project root, os.path.dirname(os.path.abspath(__file__)) is fine for its own constants.
# But if functions here (like generate_pdf_for_document) need the *application's* root directory (where main.py is, or the sys._MEIPASS dir),
# it must be passed in or imported carefully to avoid circular dependencies if utils itself is imported by main.
# For generate_pdf_for_document, I've added MAIN_APP_ROOT_DIR import as a temporary measure.
# For load_config, it now explicitly takes app_root_dir.
# populate_docx_template does not seem to need APP_ROOT_DIR directly.
# save_config and get_config_file_path/get_config_dir are self-contained.
