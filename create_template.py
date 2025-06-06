import os
from docx import Document
from docx.shared import Pt

def create_sales_contract_template():
    contract_content = """
CONTRAT DE VENTE

ENTRE LES SOUSSIGNÉS :

[Nom de la société vendeuse], société [Forme juridique] au capital de [Montant du capital social] €, immatriculée au RCS de [Ville RCS] sous le numéro [Numéro RCS], dont le siège social est situé [Adresse du siège social de la société vendeuse], représentée par [Nom du représentant légal de la société vendeuse], en qualité de [Fonction du représentant légal].

Ci-après dénommée le « Vendeur »,

ET

[Nom de la société acheteuse], société [Forme juridique] au capital de [Montant du capital social] €, immatriculée au RCS de [Ville RCS] sous le numéro [Numéro RCS], dont le siège social est situé [Adresse du siège social de la société acheteuse], représentée par [Nom du représentant légal de la société acheteuse], en qualité de [Fonction du représentant légal].

Ci-après dénommée l’ « Acheteur »,

ARTICLE 1 – OBJET DU CONTRAT
Le présent contrat a pour objet la vente par le Vendeur à l’Acheteur des biens/services suivants :
[Description détaillée des biens/services vendus, y compris les quantités, spécifications techniques, etc.]

ARTICLE 2 – PRIX
Le prix total des biens/services objet du présent contrat est fixé à [Montant total du prix] € ([Montant en lettres] euros), hors taxes.
Ce prix est ferme et définitif.
Les modalités de paiement sont les suivantes : [Modalités de paiement : acompte, solde, échéances, etc.]

ARTICLE 3 – LIVRAISON / EXÉCUTION
La livraison des biens / l’exécution des services sera effectuée au plus tard le [Date de livraison/exécution] à l’adresse suivante : [Adresse de livraison/exécution].
Les frais de livraison sont à la charge de [Vendeur/Acheteur].

ARTICLE 4 – RÉSERVE DE PROPRIÉTÉ
Le Vendeur conserve la propriété pleine et entière des biens vendus jusqu’au paiement effectif de l’intégralité du prix en principal et accessoires.

ARTICLE 5 – GARANTIES
Les biens/services vendus sont garantis contre tout vice caché ou défaut de conformité pendant une durée de [Durée de la garantie] à compter de la date de livraison/exécution.

ARTICLE 6 – RESPONSABILITÉ
La responsabilité du Vendeur ne pourra être engagée en cas de dommages indirects subis par l’Acheteur. En tout état de cause, la responsabilité du Vendeur est limitée au montant effectivement payé par l’Acheteur au titre du présent contrat.

ARTICLE 7 – FORCE MAJEURE
Aucune des parties ne pourra être tenue pour responsable des manquements ou retards dans l’exécution de ses obligations contractuelles dus à la survenance d’un cas de force majeure.

ARTICLE 8 – RÉSILIATION
En cas de manquement par l’une des parties à ses obligations contractuelles, l’autre partie pourra résilier le présent contrat de plein droit [Nombre de jours] jours après une mise en demeure restée infructueuse.

ARTICLE 9 – DROIT APPLICABLE ET LITIGES
Le présent contrat est soumis au droit français. Tout litige relatif à l’interprétation ou à l’exécution du présent contrat sera de la compétence exclusive du Tribunal de Commerce de [Ville du tribunal compétent].

Fait en deux exemplaires originaux, à [Lieu de signature], le [Date de signature].

Pour le Vendeur :
[Nom du représentant légal de la société vendeuse]
[Signature]

Pour l’Acheteur :
[Nom du représentant légal de la société acheteuse]
[Signature]
"""

    doc = Document()

    # Title
    title = doc.add_heading("CONTRAT DE VENTE", level=1)
    title.alignment = 1 # WD_ALIGN_PARAGRAPH.CENTER but using integer value

    # Add content as paragraphs, preserving line breaks roughly
    # Split content into lines and add them. Empty lines will create space.
    lines = contract_content.strip().split('\n')
    for line in lines:
        if line.strip() == "": # Add an empty paragraph for spacing if the line is empty
            doc.add_paragraph("")
        else:
            p = doc.add_paragraph()
            run = p.add_run(line)
            # Basic styling (optional, can be refined)
            if line.startswith("ARTICLE") or "ENTRE LES SOUSSIGNÉS" in line or "Pour le Vendeur" in line or "Pour l’Acheteur" in line:
                run.bold = True
            if "Ci-après dénommée" in line:
                run.italic = True
            # Default font size can be set document-wide or per run
            # font = run.font
            # font.name = 'Calibri'
            # font.size = Pt(11)

    # Ensure directory exists
    output_dir = "templates/fr/"
    os.makedirs(output_dir, exist_ok=True)

    file_path = os.path.join(output_dir, "contrat_vente_template.docx")
    doc.save(file_path)
    print(f"Document saved to {file_path}")

if __name__ == "__main__":
    create_sales_contract_template()
