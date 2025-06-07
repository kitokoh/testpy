# -*- coding: utf-8 -*-
import sys
import os
import json
import shutil
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from datetime import datetime, timedelta

from PyQt5.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QPushButton, QLabel, QLineEdit, QTextEdit,
    QListWidget, QFileDialog, QMessageBox, QDialog, QFormLayout, QComboBox,
    QDialogButtonBox, QTableWidget, QTableWidgetItem, QAbstractItemView,
    QHeaderView, QInputDialog, QGroupBox, QCheckBox, QListWidgetItem, QDoubleSpinBox,
    QStackedWidget # Added for preview
)
from PyQt5.QtGui import QIcon, QDesktopServices, QFont, QColor, QPixmap # Added QPixmap
from PyQt5.QtCore import Qt, QUrl, QDir, QCoreApplication, pyqtSignal # Added pyqtSignal
from PyQt5.QtWebEngineWidgets import QWebEngineView # Added for HTML preview

from app_config import CONFIG, APP_ROOT_DIR, load_config
import openpyxl # Added for Excel preview
from docx import Document # Added for Word preview (already here but good to note)

# Imports from other project files
import db as db_manager # Assuming db.py is accessible
from excel_editor import ExcelEditor
from html_editor import HtmlEditor
from html_to_pdf_util import render_html_template, convert_html_to_pdf # Added
from pagedegrde import generate_cover_page_logic, APP_CONFIG as PAGEDEGRDE_APP_CONFIG
from docx import Document # For populate_docx_template

# Imports from main_window.py (or a future config.py if we make one)
# These will be resolved once main_window.py is created and these vars are defined there.
# from main_window import CONFIG, APP_ROOT_DIR, DATABASE_NAME, load_config # load_config for CompilePdfDialog.send_email
# DATABASE_NAME might be better accessed via db_manager if possible, or CONFIG.
# For now, assume direct import.

# PyPDF2 for CompilePdfDialog
from PyPDF2 import PdfMerger

# Replicated Helper Styles (from main_window.py / projectManagement.py)
# These are simplified versions for use within this file.
# Ideally, these would be in a shared styling module.

STYLE_GENERIC_INPUT = """
    QLineEdit, QComboBox, QTextEdit, QDoubleSpinBox, QSpinBox, QDateEdit {
        padding: 8px 10px; border: 1px solid #ced4da;
        border-radius: 4px; background-color: white; min-height: 20px;
    }
    QLineEdit:focus, QComboBox:focus, QTextEdit:focus, QDoubleSpinBox:focus, QSpinBox:focus, QDateEdit:focus {
        border-color: #80bdff;
    }
    QComboBox::drop-down {
        subcontrol-origin: padding; subcontrol-position: top right; width: 20px;
        border-left-width: 1px; border-left-color: #ced4da; border-left-style: solid;
        border-top-right-radius: 3px; border-bottom-right-radius: 3px;
    }
    QComboBox::down-arrow { image: url(icons/arrow_down.png); }
"""

STYLE_PRIMARY_BUTTON = """
    QPushButton {
        background-color: #28a745; color: white; font-weight: bold;
        padding: 8px 12px; border-radius: 4px; border: none;
    }
    QPushButton:hover { background-color: #218838; }
    QPushButton:pressed { background-color: #1e7e34; }
"""

STYLE_SECONDARY_BUTTON = """
    QPushButton {
        background-color: #007bff; color: white; font-weight: bold;
        padding: 8px 12px; border-radius: 4px; border: none;
    }
    QPushButton:hover { background-color: #0069d9; }
    QPushButton:pressed { background-color: #005cbf; }
"""

STYLE_DANGER_BUTTON = """
    QPushButton {
        background-color: #dc3545; color: white; font-weight: bold;
        padding: 8px 12px; border-radius: 4px; border: none;
    }
    QPushButton:hover { background-color: #c82333; }
    QPushButton:pressed { background-color: #bd2130; }
"""

STYLE_NEUTRAL_BUTTON = """
    QPushButton {
        background-color: #6c757d; color: white; font-weight: normal;
        padding: 8px 12px; border-radius: 4px; border: none;
    }
    QPushButton:hover { background-color: #5a6268; }
    QPushButton:pressed { background-color: #545b62; }
"""


class ContactDialog(QDialog):
    def __init__(self, client_id=None, contact_data=None, parent=None):
        super().__init__(parent)
        # self.client_id = client_id # Not directly used if contact_data has all info
        self.contact_data = contact_data or {}
        self.setWindowTitle(self.tr("Gestion des Contacts") if not contact_data else self.tr("Modifier Contact"))
        self.setMinimumSize(400, 270) # Adjusted size
        self.setStyleSheet(STYLE_GENERIC_INPUT) # Apply to all inputs in dialog
        self.setup_ui()

    def setup_ui(self):
        layout = QFormLayout(self)
        self.name_input = QLineEdit(self.contact_data.get("name", ""))
        layout.addRow(self.tr("Nom complet:"), self.name_input)
        self.email_input = QLineEdit(self.contact_data.get("email", ""))
        layout.addRow(self.tr("Email:"), self.email_input)
        self.phone_input = QLineEdit(self.contact_data.get("phone", ""))
        layout.addRow(self.tr("Téléphone:"), self.phone_input)
        self.position_input = QLineEdit(self.contact_data.get("position", ""))
        layout.addRow(self.tr("Poste:"), self.position_input)
        self.primary_check = QCheckBox(self.tr("Contact principal"))
        self.primary_check.setChecked(bool(self.contact_data.get("is_primary", 0)))
        layout.addRow(self.primary_check)

        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        ok_button = button_box.button(QDialogButtonBox.Ok)
        ok_button.setText(self.tr("OK"))
        ok_button.setStyleSheet(STYLE_PRIMARY_BUTTON)

        cancel_button = button_box.button(QDialogButtonBox.Cancel)
        cancel_button.setText(self.tr("Annuler"))
        # cancel_button.setStyleSheet(STYLE_NEUTRAL_BUTTON) # Optional: for explicit neutral

        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        layout.addRow(button_box)

    def get_data(self):
        return {
            "name": self.name_input.text().strip(),
            "email": self.email_input.text().strip(),
            "phone": self.phone_input.text().strip(),
            "position": self.position_input.text().strip(),
            "is_primary": 1 if self.primary_check.isChecked() else 0
        }

class TemplateDialog(QDialog):
    def __init__(self, parent=None): # Removed config from constructor, will use global CONFIG
        super().__init__(parent)
        self.setWindowTitle(self.tr("Select Template"))
        self.setMinimumSize(850, 650) # Slightly increased dialog size for comfort

        # Initialize UI elements for preview
        self.preview_stacked_widget = QStackedWidget()
        self.default_preview_label = QLabel(self.tr("Select a template from the list to see its preview here.")) # More informative text
        self.default_preview_label.setAlignment(Qt.AlignCenter)
        self.default_preview_label.setStyleSheet("""
            QLabel {
                font-size: 14px;
                color: #888;
                border: 1px dashed #ccc;
                border-radius: 5px;
                padding: 20px;
            }
        """)
        self.image_text_preview_label = QLabel()
        self.image_text_preview_label.setAlignment(Qt.AlignCenter) # Default, will be overridden for text
        self.html_preview_widget = QWebEngineView()

        self.preview_stacked_widget.addWidget(self.default_preview_label)
        self.preview_stacked_widget.addWidget(self.image_text_preview_label)
        self.preview_stacked_widget.addWidget(self.html_preview_widget)

        self.setup_ui()

    def setup_ui(self):
        main_layout = QHBoxLayout(self)
        main_layout.setSpacing(10) # Add some spacing between splitter panes

        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        left_layout.setContentsMargins(0,0,0,0) # No margins for the layout itself

        left_layout.addWidget(QLabel(self.tr("Available Templates:")))
        self.template_list = QListWidget()
        self.template_list.itemDoubleClicked.connect(self.edit_template)
        self.template_list.itemClicked.connect(self.display_template_preview)
        self.template_list.setSpacing(3) # Add spacing between list items
        left_layout.addWidget(self.template_list)

        template_actions_layout = QGridLayout()
        template_actions_layout.setSpacing(5)
        self.add_btn = QPushButton(self.tr("Add"))
        self.add_btn.setIcon(QIcon.fromTheme("list-add"))
        self.add_btn.clicked.connect(self.add_template)
        template_actions_layout.addWidget(self.add_btn, 0, 0)

        self.edit_btn = QPushButton(self.tr("Edit Path")) # Clarified Edit action
        self.edit_btn.setIcon(QIcon.fromTheme("document-edit"))
        self.edit_btn.clicked.connect(self.edit_template) # This opens the file, consider renaming if behavior changes
        template_actions_layout.addWidget(self.edit_btn, 0, 1)

        self.delete_btn = QPushButton(self.tr("Delete"))
        self.delete_btn.setIcon(QIcon.fromTheme("edit-delete"))
        self.delete_btn.clicked.connect(self.delete_template)
        template_actions_layout.addWidget(self.delete_btn, 1, 0)

        self.default_btn = QPushButton(self.tr("Set Default"))
        self.default_btn.setIcon(QIcon.fromTheme("emblem-default"))
        self.default_btn.clicked.connect(self.set_default_template)
        template_actions_layout.addWidget(self.default_btn, 1, 1)
        left_layout.addLayout(template_actions_layout)

        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        right_layout.setContentsMargins(0,0,0,0)

        preview_label_title = QLabel(self.tr("Preview:"))
        preview_label_title.setFont(QFont("Arial", 10, QFont.Bold)) # Slightly bolder title for preview
        right_layout.addWidget(preview_label_title, 0)
        right_layout.addWidget(self.preview_stacked_widget, 1)

        self.button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        self.button_box.button(QDialogButtonBox.Ok).setText(self.tr("OK"))
        self.button_box.button(QDialogButtonBox.Ok).setIcon(QIcon.fromTheme("dialog-ok-apply"))
        self.button_box.button(QDialogButtonBox.Cancel).setText(self.tr("Cancel"))
        self.button_box.button(QDialogButtonBox.Cancel).setIcon(QIcon.fromTheme("dialog-cancel"))
        self.button_box.accepted.connect(self.accept)
        self.button_box.rejected.connect(self.reject)
        right_layout.addWidget(self.button_box)

        splitter = QSplitter(Qt.Horizontal)
        splitter.addWidget(left_widget)
        splitter.addWidget(right_widget)
        splitter.setSizes([320, 530]) # Adjusted proportions slightly
        splitter.setStretchFactor(0, 3) # Left panel less stretchy
        splitter.setStretchFactor(1, 7) # Right panel more stretchy


        main_layout.addWidget(splitter)
        self.setLayout(main_layout)

        self.load_templates()
        self.preview_stacked_widget.setCurrentWidget(self.default_preview_label)

    def load_templates(self):
        self.template_list.clear()
        try:
            all_templates = db_manager.get_all_templates() # Changed to get_all_templates
            if all_templates is None: all_templates = []
            for template_data in all_templates:
                item_text = f"{template_data.get('template_name', 'N/A')} ({template_data.get('language_code', 'N/A')})"
                if template_data.get('is_default_for_type_lang'): # Changed to is_default_for_type_lang
                    item_text += f" [{self.tr('Default')}]"

                item = QListWidgetItem(item_text)
                full_path = os.path.join(CONFIG["templates_dir"], template_data.get('language_code', ''), template_data.get('base_file_name', ''))

                # Set Icon based on file type
                file_ext = os.path.splitext(full_path)[1].lower()
                icon_name = "text-x-generic" # Default icon
                if file_ext == '.xlsx':
                    icon_name = "application-vnd.ms-excel" # Common theme name for Excel
                elif file_ext == '.docx':
                    icon_name = "application-msword" # Common theme name for Word
                elif file_ext in ['.html', '.htm']:
                    icon_name = "text-html"
                elif file_ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']:
                    icon_name = "image-x-generic"

                item.setIcon(QIcon.fromTheme(icon_name, QIcon.fromTheme("text-x-generic"))) # Fallback to generic text icon

                item.setData(Qt.UserRole, full_path)
                item.setData(Qt.UserRole + 1, template_data.get('template_id'))
                self.template_list.addItem(item)
        except Exception as e:
            QMessageBox.warning(self, self.tr("Database Error"), self.tr("Error loading templates:
{0}").format(str(e)))

    def add_template(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, self.tr("Sélectionner un modèle"), CONFIG["templates_dir"],
            self.tr("Fichiers Modèles (*.xlsx *.docx *.html);;Fichiers Excel (*.xlsx);;Documents Word (*.docx);;Documents HTML (*.html);;Tous les fichiers (*)")
        )
        if not file_path: return
        name, ok = QInputDialog.getText(self, self.tr("Nom du Modèle"), self.tr("Entrez un nom pour ce modèle:"))
        if not ok or not name.strip(): return
        languages = ["fr", "en", "ar", "tr", "pt"] # These should ideally come from a central config or db
        lang, ok = QInputDialog.getItem(self, self.tr("Langue du Modèle"), self.tr("Sélectionnez la langue:"), languages, 0, False)
        if not ok: return

        target_dir = os.path.join(CONFIG["templates_dir"], lang)
        os.makedirs(target_dir, exist_ok=True)
        base_file_name = os.path.basename(file_path)
        target_path = os.path.join(target_dir, base_file_name)
        try:
            shutil.copy(file_path, target_path)
            # db_manager.add_template should handle the DB insertion
            # It needs: name, base_file_name, lang, type (Excel, Word, HTML), is_default (false initially)
            file_ext = os.path.splitext(base_file_name)[1].lower()
            doc_type = "Excel" # Default
            if file_ext == ".docx": doc_type = "Word"
            elif file_ext == ".html": doc_type = "HTML"

            template_data = {
                'template_name': name.strip(),
                'base_file_name': base_file_name,
                'language_code': lang,
                'document_type': doc_type,
                'description': '', # Optional: add a field for description
                'is_default_for_category_language': False
            }
            new_id = db_manager.add_template(template_data) # Assumes this function exists
            if new_id:
                self.load_templates()
                QMessageBox.information(self, self.tr("Succès"), self.tr("Modèle ajouté avec succès."))
            else:
                QMessageBox.critical(self, self.tr("Erreur"), self.tr("Erreur lors de l'ajout du modèle à la DB."))
        except Exception as e:
            QMessageBox.critical(self, self.tr("Erreur"), self.tr("Erreur lors de l'ajout du modèle:
{0}").format(str(e)))

    def edit_template(self):
        item = self.template_list.currentItem()
        if not item: return
        # Edit now refers to opening the file path, which is stored in UserRole
        file_path = item.data(Qt.UserRole)
        if file_path and os.path.exists(file_path):
            QDesktopServices.openUrl(QUrl.fromLocalFile(file_path))
        else:
            # Fallback or error if path is not valid, possibly using template_id for refetch if needed
            template_id = item.data(Qt.UserRole + 1)
            try:
                template_details = db_manager.get_template_by_id(template_id)
                if template_details:
                    correct_file_path = os.path.join(CONFIG["templates_dir"], template_details.get('language_code'), template_details.get('base_file_name'))
                    if os.path.exists(correct_file_path):
                         QDesktopServices.openUrl(QUrl.fromLocalFile(correct_file_path))
                    else:
                         QMessageBox.warning(self, self.tr("File Not Found"), self.tr("Template file not found at: {0}").format(correct_file_path))
                else:
                    QMessageBox.warning(self, self.tr("Error"), self.tr("Template details not found in DB."))
            except Exception as e:
                QMessageBox.warning(self, self.tr("Database Error"), self.tr("Error accessing template details:
{0}").format(str(e)))


    def delete_template(self):
        item = self.template_list.currentItem()
        if not item: return
        template_id = item.data(Qt.UserRole + 1) # ID is in UserRole + 1
        reply = QMessageBox.question(self, self.tr("Confirm Deletion"), self.tr("Are you sure you want to delete this template?"), QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
        if reply == QMessageBox.Yes:
            try:
                template_details = db_manager.get_template_by_id(template_id)
                deleted_from_db = db_manager.delete_template(template_id)
                if deleted_from_db and template_details:
                    file_path_to_delete = os.path.join(CONFIG["templates_dir"], template_details.get('language_code'), template_details.get('base_file_name'))
                    if os.path.exists(file_path_to_delete): os.remove(file_path_to_delete)
                    self.load_templates()
                    self.preview_stacked_widget.setCurrentWidget(self.default_preview_label) # Reset preview
                    QMessageBox.information(self, self.tr("Success"), self.tr("Template deleted successfully."))
                elif not template_details:
                     QMessageBox.warning(self, self.tr("Error"), self.tr("Template not found in DB."))
                else:
                     QMessageBox.critical(self, self.tr("Error"), self.tr("Error deleting template from DB."))
            except Exception as e:
                QMessageBox.critical(self, self.tr("Error"), self.tr("Error deleting template:
{0}").format(str(e)))

    def set_default_template(self):
        item = self.template_list.currentItem()
        if not item: return
        template_id = item.data(Qt.UserRole + 1) # ID is in UserRole + 1
        try:
            # Ensure the correct function name is called here if it changed in db_manager
            # For now, assuming set_default_template_by_id is the correct one for setting default status
            success = db_manager.set_default_template_by_id(template_id)
            if success:
                self.load_templates()
                QMessageBox.information(self, self.tr("Success"), self.tr("Template set as default for its category and language."))
            else:
                QMessageBox.critical(self, self.tr("Database Error"), self.tr("Error setting default template."))
        except Exception as e:
            QMessageBox.critical(self, self.tr("Database Error"), self.tr("Error updating template:
{0}").format(str(e)))

    def get_selected_template_path(self):
        selected_items = self.template_list.selectedItems()
        if selected_items:
            return selected_items[0].data(Qt.UserRole) # File path is in Qt.UserRole
        return None

    def display_template_preview(self, item):
        file_path = item.data(Qt.UserRole)

        # Clear previous pixmap if any, before loading new content or showing text
        self.image_text_preview_label.setPixmap(QPixmap()) # Clear any existing image

        if not file_path or not os.path.exists(file_path):
            self.default_preview_label.setText(self.tr("File not found: {0}").format(os.path.basename(file_path) if file_path else "N/A"))
            self.preview_stacked_widget.setCurrentWidget(self.default_preview_label)
            return

        self.image_text_preview_label.setAlignment(Qt.AlignCenter)
        self.image_text_preview_label.setWordWrap(False)

        try:
            ext = os.path.splitext(file_path)[1].lower()
            filename = os.path.basename(file_path)

            if ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']:
                pixmap = QPixmap(file_path)
                if pixmap.isNull():
                    self.image_text_preview_label.setText(self.tr("Error: Could not load image '{0}'.").format(filename))
                else:
                    target_size = self.image_text_preview_label.size()
                    if not target_size.isValid() or target_size.width() < 50 or target_size.height() < 50 :
                        target_size = self.preview_stacked_widget.size()

                    scaled_pixmap = pixmap.scaled(target_size.width() - 10 , target_size.height() - 10,
                                                  Qt.KeepAspectRatio, Qt.SmoothTransformation)
                    self.image_text_preview_label.setPixmap(scaled_pixmap)
                self.preview_stacked_widget.setCurrentWidget(self.image_text_preview_label)
            elif ext == '.xlsx':
                try:
                    workbook = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
                    sheet = workbook.active
                    preview_text = f"<b>{filename}</b>\n\n"
                    max_rows_preview = 15
                    max_cols_preview = 7
                    for row_idx, row_data in enumerate(sheet.iter_rows(max_row=max_rows_preview, max_col=max_cols_preview)):
                        preview_text += "\t".join([str(cell.value) if cell.value is not None else "" for cell in row_data]) + "\n"

                    if sheet.max_row > max_rows_preview or sheet.max_column > max_cols_preview:
                        preview_text += self.tr("\n[Preview truncated]...")

                    self.image_text_preview_label.setText(preview_text.strip())
                    self.image_text_preview_label.setAlignment(Qt.AlignTop | Qt.AlignLeft)
                    self.image_text_preview_label.setWordWrap(True)
                except Exception as e:
                    self.image_text_preview_label.setText(self.tr("Could not generate preview for '{0}'.\nError: {1}").format(filename, str(e)))
                self.preview_stacked_widget.setCurrentWidget(self.image_text_preview_label)
            elif ext == '.docx':
                try:
                    doc = Document(file_path)
                    preview_text = f"<b>{filename}</b>\n\n"
                    max_paras_preview = 7
                    preview_text += "\n\n".join([para.text for para in doc.paragraphs[:max_paras_preview]]).strip()

                    if len(doc.paragraphs) > max_paras_preview:
                        preview_text += self.tr("\n\n[Preview truncated]...")

                    if not preview_text.strip() == f"<b>{filename}</b>": # check if any content was added
                         self.image_text_preview_label.setText(preview_text)
                    else: # Only filename was added, means no text in first paras
                         self.image_text_preview_label.setText(f"<b>{filename}</b>\n\n{self.tr('[Empty Document or No Text in First {0} Paragraphs]').format(max_paras_preview)}")

                    self.image_text_preview_label.setAlignment(Qt.AlignTop | Qt.AlignLeft)
                    self.image_text_preview_label.setWordWrap(True)
                except Exception as e:
                    self.image_text_preview_label.setText(self.tr("Could not generate preview for '{0}'.\nError: {1}").format(filename, str(e)))
                self.preview_stacked_widget.setCurrentWidget(self.image_text_preview_label)
            elif ext == '.html' or ext == '.htm':
                self.html_preview_widget.setUrl(QUrl.fromLocalFile(file_path))
                self.preview_stacked_widget.setCurrentWidget(self.html_preview_widget)
            elif ext in ['.txt', '.md', '.csv', '.json', '.xml', '.py', '.js', '.css']:
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        file_content = f.read(4096)
                        preview_text = f"<b>{filename}</b>\n\n{file_content}"
                        if len(file_content) == 4096:
                            preview_text += self.tr("\n\n[Preview truncated]...")
                    self.image_text_preview_label.setText(preview_text.strip())
                    self.image_text_preview_label.setAlignment(Qt.AlignTop | Qt.AlignLeft)
                    self.image_text_preview_label.setWordWrap(True)
                except Exception as e:
                    self.image_text_preview_label.setText(self.tr("Could not generate preview for '{0}'.\nError: {1}").format(filename, str(e)))
                self.preview_stacked_widget.setCurrentWidget(self.image_text_preview_label)
            else:
                self.image_text_preview_label.setText(self.tr("No preview available for '{0}' ({1}).").format(filename, ext))
                self.image_text_preview_label.setAlignment(Qt.AlignCenter)
                self.preview_stacked_widget.setCurrentWidget(self.image_text_preview_label)
        except Exception as e:
            self.default_preview_label.setText(self.tr("Error generating preview for {0}.\nDetails: {1}").format(os.path.basename(file_path) if file_path else "N/A", str(e)))
            self.preview_stacked_widget.setCurrentWidget(self.default_preview_label)

class ProductDialog(QDialog):
    def __init__(self, client_id, product_data=None, parent=None): # client_id might not be needed if product_data is self-sufficient for editing
        super().__init__(parent)
        # self.client_id = client_id # Store if needed for linking new products
        self.product_data = product_data or {} # For editing existing linked product
        self.setWindowTitle(self.tr("Ajouter Produit") if not self.product_data.get('client_project_product_id') else self.tr("Modifier Produit"))
        self.setStyleSheet(STYLE_GENERIC_INPUT)
        self.setup_ui()

    def setup_ui(self):
        layout = QFormLayout(self)
        self.name_input = QLineEdit(self.product_data.get("product_name", "")) # Global product name
        layout.addRow(self.tr("Nom du Produit:"), self.name_input)
        self.description_input = QTextEdit(self.product_data.get("product_description", "")) # Global product desc
        self.description_input.setFixedHeight(80)
        layout.addRow(self.tr("Description:"), self.description_input)
        self.quantity_input = QDoubleSpinBox()
        self.quantity_input.setRange(0, 1000000)
        self.quantity_input.setValue(self.product_data.get("quantity", 0)) # From ClientProjectProducts
        self.quantity_input.valueChanged.connect(self.update_total_price)
        layout.addRow(self.tr("Quantité:"), self.quantity_input)

        # Effective unit price (override or base)
        effective_unit_price = self.product_data.get('unit_price_override', self.product_data.get('base_unit_price', 0.0))
        self.unit_price_input = QDoubleSpinBox()
        self.unit_price_input.setRange(0, 10000000)
        self.unit_price_input.setPrefix("€ ")
        self.unit_price_input.setValue(effective_unit_price)
        self.unit_price_input.valueChanged.connect(self.update_total_price)
        layout.addRow(self.tr("Prix Unitaire:"), self.unit_price_input)

        total_price_title_label = QLabel(self.tr("Prix Total:"))
        self.total_price_label = QLabel("€ 0.00")
        self.total_price_label.setStyleSheet("font-weight: bold;")
        layout.addRow(total_price_title_label, self.total_price_label)
        self.update_total_price() # Initial calculation

        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        ok_button = button_box.button(QDialogButtonBox.Ok)
        ok_button.setText(self.tr("OK"))
        ok_button.setStyleSheet(STYLE_PRIMARY_BUTTON)
        cancel_button = button_box.button(QDialogButtonBox.Cancel)
        cancel_button.setText(self.tr("Annuler"))
        # cancel_button.setStyleSheet(STYLE_NEUTRAL_BUTTON)

        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        layout.addRow(button_box)

    def update_total_price(self):
        quantity = self.quantity_input.value()
        unit_price = self.unit_price_input.value()
        total = quantity * unit_price
        self.total_price_label.setText(f"€ {total:.2f}")

    def get_data(self):
        # Returns data for both global product and the link
        return {
            # Global product fields (for add/update global product)
            "product_name": self.name_input.text().strip(),
            "product_description": self.description_input.toPlainText().strip(),
            # Link fields (for ClientProjectProducts)
            "quantity": self.quantity_input.value(),
            "unit_price_for_dialog": self.unit_price_input.value(), # This is the price user set in dialog
            # "total_price": self.quantity_input.value() * self.unit_price_input.value() # Calculated, not stored directly here
        }

class CreateDocumentDialog(QDialog):
    def __init__(self, client_info, parent=None): # Removed config, uses global CONFIG
        super().__init__(parent)
        self.client_info = client_info
        self.setWindowTitle(self.tr("Créer des Documents"))
        self.setMinimumSize(600, 400)
        self.setup_ui()

    def setup_ui(self):
        layout = QVBoxLayout(self)
        lang_layout = QHBoxLayout()
        lang_layout.addWidget(QLabel(self.tr("Langue:")))
        self.lang_combo = QComboBox()
        self.lang_combo.addItems(self.client_info.get("selected_languages", ["fr"]))
        lang_layout.addWidget(self.lang_combo)
        layout.addLayout(lang_layout)
        self.templates_list = QListWidget()
        self.templates_list.setSelectionMode(QListWidget.MultiSelection)
        layout.addWidget(QLabel(self.tr("Sélectionnez les documents à créer:")))
        layout.addWidget(self.templates_list)
        self.load_templates_for_creation() # Renamed to avoid conflict
        btn_layout = QHBoxLayout()
        create_btn = QPushButton(self.tr("Créer Documents"))
        create_btn.setIcon(QIcon.fromTheme("document-new"))
        create_btn.clicked.connect(self.create_documents_from_selection) # Renamed
        btn_layout.addWidget(create_btn)
        cancel_btn = QPushButton(self.tr("Annuler"))
        cancel_btn.setIcon(QIcon.fromTheme("dialog-cancel"))
        cancel_btn.clicked.connect(self.reject)
        btn_layout.addWidget(cancel_btn)
        layout.addLayout(btn_layout)

    def load_templates_for_creation(self):
        self.templates_list.clear()
        try:
            # Fetch templates suitable for document generation (e.g., those with base_file_name)
            # db_manager.get_all_file_based_templates() or similar is needed
            all_file_templates = db_manager.get_all_file_based_templates() # Assumed function
            if all_file_templates is None: all_file_templates = []
            for template_dict in all_file_templates:
                name = template_dict.get('template_name', 'N/A')
                lang = template_dict.get('language_code', 'N/A')
                base_file_name = template_dict.get('base_file_name', 'N/A')
                item_text = f"{name} ({lang}) - {base_file_name}"
                item = QListWidgetItem(item_text)
                item.setData(Qt.UserRole, (name, lang, base_file_name)) # Store data for creation
                self.templates_list.addItem(item)
        except Exception as e:
            QMessageBox.warning(self, self.tr("Erreur DB"), self.tr("Erreur de chargement des modèles:
{0}").format(str(e)))

    def create_documents_from_selection(self): # Renamed
        selected_items = self.templates_list.selectedItems()
        if not selected_items:
            QMessageBox.warning(self, self.tr("Aucun document sélectionné"), self.tr("Veuillez sélectionner au moins un document à créer."))
            return
        lang = self.lang_combo.currentText()
        target_dir = os.path.join(self.client_info["base_folder_path"], lang)
        os.makedirs(target_dir, exist_ok=True)
        created_files = []
        for item in selected_items:
            db_template_name, db_template_lang, actual_template_filename = item.data(Qt.UserRole)
            if not actual_template_filename:
                QMessageBox.warning(self, self.tr("Erreur Modèle"), self.tr("Nom de fichier manquant pour le modèle '{0}'. Impossible de créer.").format(db_template_name))
                continue
            template_file_found_abs = os.path.join(CONFIG["templates_dir"], db_template_lang, actual_template_filename)
            if os.path.exists(template_file_found_abs):
                target_path = os.path.join(target_dir, actual_template_filename)
                shutil.copy(template_file_found_abs, target_path)
                if target_path.lower().endswith(".docx"):
                    try:
                        populate_docx_template(target_path, self.client_info) # Defined below
                    except Exception as e_pop:
                        QMessageBox.warning(self, self.tr("Erreur DOCX"), self.tr("Impossible de populer le modèle Word '{0}':
{1}").format(os.path.basename(target_path), e_pop))
                elif target_path.lower().endswith(".html"):
                    try:
                        with open(target_path, 'r', encoding='utf-8') as f: template_content = f.read()
                        populated_content = HtmlEditor.populate_html_content(template_content, self.client_info)
                        with open(target_path, 'w', encoding='utf-8') as f: f.write(populated_content)
                    except Exception as e_html_pop:
                        QMessageBox.warning(self, self.tr("Erreur HTML"), self.tr("Impossible de populer le modèle HTML '{0}':
{1}").format(os.path.basename(target_path), e_html_pop))
                created_files.append(target_path)
            else:
                QMessageBox.warning(self, self.tr("Erreur Modèle"), self.tr("Fichier modèle '{0}' introuvable pour '{1}'.").format(actual_template_filename, db_template_name))
        if created_files:
            QMessageBox.information(self, self.tr("Documents créés"), self.tr("{0} documents ont été créés avec succès.").format(len(created_files)))
            self.accept()
        else:
            QMessageBox.warning(self, self.tr("Erreur"), self.tr("Aucun document n'a pu être créé."))

class CompilePdfDialog(QDialog):
    def __init__(self, client_info, parent=None):
        super().__init__(parent)
        self.client_info = client_info
        self.setWindowTitle(self.tr("Compiler des PDF"))
        self.setMinimumSize(700, 500)
        self.setup_ui()

    def setup_ui(self):
        layout = QVBoxLayout(self)
        layout.addWidget(QLabel(self.tr("Sélectionnez les PDF à compiler:")))
        self.pdf_list = QTableWidget()
        self.pdf_list.setColumnCount(4)
        self.pdf_list.setHorizontalHeaderLabels([self.tr("Sélection"), self.tr("Nom du fichier"), self.tr("Chemin"), self.tr("Pages (ex: 1-3,5)")])
        self.pdf_list.horizontalHeader().setSectionResizeMode(1, QHeaderView.Stretch)
        self.pdf_list.setSelectionBehavior(QAbstractItemView.SelectRows)
        layout.addWidget(self.pdf_list)
        btn_layout = QHBoxLayout()
        add_btn = QPushButton(self.tr("Ajouter PDF")); add_btn.setIcon(QIcon.fromTheme("list-add")); add_btn.clicked.connect(self.add_pdf_to_list); btn_layout.addWidget(add_btn)
        remove_btn = QPushButton(self.tr("Supprimer")); remove_btn.setIcon(QIcon.fromTheme("edit-delete")); remove_btn.clicked.connect(self.remove_selected_pdf); btn_layout.addWidget(remove_btn)
        move_up_btn = QPushButton(self.tr("Monter")); move_up_btn.setIcon(QIcon.fromTheme("go-up")); move_up_btn.clicked.connect(self.move_pdf_up); btn_layout.addWidget(move_up_btn)
        move_down_btn = QPushButton(self.tr("Descendre")); move_down_btn.setIcon(QIcon.fromTheme("go-down")); move_down_btn.clicked.connect(self.move_pdf_down); btn_layout.addWidget(move_down_btn)
        layout.addLayout(btn_layout)
        options_layout = QHBoxLayout()
        options_layout.addWidget(QLabel(self.tr("Nom du fichier compilé:")))
        self.output_name_edit = QLineEdit(f"{self.tr('compilation')}_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf") # Renamed
        options_layout.addWidget(self.output_name_edit)
        layout.addLayout(options_layout)
        action_layout = QHBoxLayout()
        compile_btn = QPushButton(self.tr("Compiler PDF")); compile_btn.setIcon(QIcon.fromTheme("document-export")); compile_btn.setStyleSheet("background-color: #27ae60; color: white; font-weight: bold;"); compile_btn.clicked.connect(self.execute_compile_pdf); action_layout.addWidget(compile_btn)
        cancel_btn = QPushButton(self.tr("Annuler")); cancel_btn.setIcon(QIcon.fromTheme("dialog-cancel")); cancel_btn.clicked.connect(self.reject); action_layout.addWidget(cancel_btn)
        layout.addLayout(action_layout)
        self.load_existing_pdfs_for_client() # Renamed

    def load_existing_pdfs_for_client(self): # Renamed
        client_dir = self.client_info["base_folder_path"]
        pdf_files = []
        for root, _, files in os.walk(client_dir):
            for file in files:
                if file.lower().endswith('.pdf'): pdf_files.append(os.path.join(root, file))
        self.pdf_list.setRowCount(len(pdf_files))
        for i, file_path in enumerate(pdf_files):
            chk = QCheckBox(); chk.setChecked(True); self.pdf_list.setCellWidget(i, 0, chk)
            self.pdf_list.setItem(i, 1, QTableWidgetItem(os.path.basename(file_path)))
            self.pdf_list.setItem(i, 2, QTableWidgetItem(file_path))
            pages_edit = QLineEdit("all"); pages_edit.setPlaceholderText(self.tr("all ou 1-3,5")); self.pdf_list.setCellWidget(i, 3, pages_edit)

    def add_pdf_to_list(self): # Renamed
        file_paths, _ = QFileDialog.getOpenFileNames(self, self.tr("Sélectionner des PDF"), "", self.tr("Fichiers PDF (*.pdf)"))
        if not file_paths: return
        current_row_count = self.pdf_list.rowCount()
        self.pdf_list.setRowCount(current_row_count + len(file_paths))
        for i, file_path in enumerate(file_paths):
            row = current_row_count + i
            chk = QCheckBox(); chk.setChecked(True); self.pdf_list.setCellWidget(row, 0, chk)
            self.pdf_list.setItem(row, 1, QTableWidgetItem(os.path.basename(file_path)))
            self.pdf_list.setItem(row, 2, QTableWidgetItem(file_path))
            pages_edit = QLineEdit("all"); pages_edit.setPlaceholderText(self.tr("all ou 1-3,5")); self.pdf_list.setCellWidget(row, 3, pages_edit)

    def remove_selected_pdf(self): # Renamed
        selected_rows = set(index.row() for index in self.pdf_list.selectedIndexes())
        for row in sorted(selected_rows, reverse=True): self.pdf_list.removeRow(row)

    def move_pdf_up(self): # Renamed
        current_row = self.pdf_list.currentRow()
        if current_row > 0: self.swap_pdf_rows(current_row, current_row - 1); self.pdf_list.setCurrentCell(current_row - 1, 0) # Renamed

    def move_pdf_down(self): # Renamed
        current_row = self.pdf_list.currentRow()
        if current_row < self.pdf_list.rowCount() - 1: self.swap_pdf_rows(current_row, current_row + 1); self.pdf_list.setCurrentCell(current_row + 1, 0) # Renamed

    def swap_pdf_rows(self, row1, row2): # Renamed
        for col in range(self.pdf_list.columnCount()):
            item1 = self.pdf_list.takeItem(row1, col); item2 = self.pdf_list.takeItem(row2, col)
            self.pdf_list.setItem(row1, col, item2); self.pdf_list.setItem(row2, col, item1)
        widget1_chk = self.pdf_list.cellWidget(row1, 0); widget1_pages = self.pdf_list.cellWidget(row1, 3)
        widget2_chk = self.pdf_list.cellWidget(row2, 0); widget2_pages = self.pdf_list.cellWidget(row2, 3)
        self.pdf_list.setCellWidget(row1, 0, widget2_chk); self.pdf_list.setCellWidget(row1, 3, widget2_pages)
        self.pdf_list.setCellWidget(row2, 0, widget1_chk); self.pdf_list.setCellWidget(row2, 3, widget1_pages)

    def execute_compile_pdf(self): # Renamed
        merger = PdfMerger()
        output_name = self.output_name_edit.text().strip()
        if not output_name: QMessageBox.warning(self, self.tr("Nom manquant"), self.tr("Veuillez spécifier un nom de fichier pour la compilation.")); return
        if not output_name.lower().endswith('.pdf'): output_name += '.pdf'
        output_path = os.path.join(self.client_info["base_folder_path"], output_name)
        cover_path = self.generate_cover_page_pdf() # Renamed
        if cover_path: merger.append(cover_path)
        for row in range(self.pdf_list.rowCount()):
            chk = self.pdf_list.cellWidget(row, 0)
            if chk and chk.isChecked():
                file_path = self.pdf_list.item(row, 2).text()
                pages_spec = self.pdf_list.cellWidget(row, 3).text().strip()
                try:
                    if pages_spec.lower() == "all" or not pages_spec: merger.append(file_path)
                    else:
                        pages = []
                        for part in pages_spec.split(','):
                            if '-' in part: start, end = part.split('-'); pages.extend(range(int(start), int(end)+1))
                            else: pages.append(int(part))
                        merger.append(file_path, pages=[p-1 for p in pages])
                except Exception as e: QMessageBox.warning(self, self.tr("Erreur"), self.tr("Erreur lors de l'ajout de {0}:
{1}").format(os.path.basename(file_path), str(e)))
        try:
            with open(output_path, 'wb') as f: merger.write(f)
            if cover_path and os.path.exists(cover_path): os.remove(cover_path)
            self.offer_download_or_email_compiled_pdf(output_path) # Renamed
            self.accept()
        except Exception as e: QMessageBox.critical(self, self.tr("Erreur"), self.tr("Erreur lors de la compilation du PDF:
{0}").format(str(e)))

    def generate_cover_page_pdf(self): # Renamed
        config_dict = {
            'title': self.tr("Compilation de Documents - Projet: {0}").format(self.client_info.get('project_identifier', self.tr('N/A'))),
            'subtitle': self.tr("Client: {0}").format(self.client_info.get('client_name', self.tr('N/A'))),
            'author': self.client_info.get('company_name', PAGEDEGRDE_APP_CONFIG.get('default_institution', self.tr('Votre Entreprise'))),
            'date': datetime.now().strftime('%d/%m/%Y %H:%M'), 'doc_type': self.tr("Compilation de Documents"),
            # ... other fields from original ...
            'font_name': PAGEDEGRDE_APP_CONFIG.get('default_font', 'Helvetica'), 'logo_data': None,
             'footer_text': self.tr("Document compilé le {0}").format(datetime.now().strftime('%d/%m/%Y'))
        }
        logo_path = os.path.join(APP_ROOT_DIR, "logo.png") # APP_ROOT_DIR from main_window
        if os.path.exists(logo_path):
            try:
                with open(logo_path, "rb") as f_logo: config_dict['logo_data'] = f_logo.read()
            except Exception as e_logo: print(self.tr("Erreur chargement logo.png: {0}").format(e_logo))
        try:
            pdf_bytes = generate_cover_page_logic(config_dict)
            base_temp_dir = self.client_info.get("base_folder_path", QDir.tempPath()) # QDir from QtCore
            temp_cover_filename = f"cover_page_generated_{datetime.now().strftime('%Y%m%d%H%M%S%f')}.pdf"
            temp_cover_path = os.path.join(base_temp_dir, temp_cover_filename)
            with open(temp_cover_path, "wb") as f: f.write(pdf_bytes)
            return temp_cover_path
        except Exception as e:
            QMessageBox.warning(self, self.tr("Erreur Page de Garde"), self.tr("Impossible de générer la page de garde: {0}").format(e)) # Simplified message
            return None

    def offer_download_or_email_compiled_pdf(self, pdf_path): # Renamed
        msg_box = QMessageBox(self); msg_box.setWindowTitle(self.tr("Compilation réussie"))
        msg_box.setText(self.tr("Le PDF compilé a été sauvegardé dans:
{0}").format(pdf_path))
        download_btn = msg_box.addButton(self.tr("Télécharger"), QMessageBox.ActionRole)
        email_btn = msg_box.addButton(self.tr("Envoyer par email"), QMessageBox.ActionRole)
        msg_box.addButton(self.tr("Fermer"), QMessageBox.RejectRole)
        msg_box.exec_()
        if msg_box.clickedButton() == download_btn: QDesktopServices.openUrl(QUrl.fromLocalFile(pdf_path))
        elif msg_box.clickedButton() == email_btn: self.send_compiled_pdf_by_email(pdf_path) # Renamed

    def send_compiled_pdf_by_email(self, pdf_path): # Renamed
        primary_email = None
        client_uuid = self.client_info.get("client_id")
        if client_uuid:
            contacts_for_client = db_manager.get_contacts_for_client(client_uuid)
            if contacts_for_client:
                for contact in contacts_for_client:
                    if contact.get('is_primary_for_client'): primary_email = contact.get('email'); break
        email, ok = QInputDialog.getText(self, self.tr("Envoyer par email"), self.tr("Adresse email du destinataire:"), text=primary_email or "")
        if not ok or not email.strip(): return
        # Uses global load_config() which is a placeholder here, will be imported from main_window
        # This is a known issue to be resolved by import adjustments.
        current_app_config = load_config()
        if not current_app_config.get("smtp_server") or not current_app_config.get("smtp_user"):
            QMessageBox.warning(self, self.tr("Configuration manquante"), self.tr("Veuillez configurer les paramètres SMTP.")) # Simplified
            return
        msg = MIMEMultipart(); msg['From'] = current_app_config["smtp_user"]; msg['To'] = email
        msg['Subject'] = self.tr("Documents compilés - {0}").format(self.client_info['client_name'])
        body = self.tr("Bonjour,

Veuillez trouver ci-joint les documents compilés pour le projet {0}.

Cordialement,
Votre équipe").format(self.client_info.get('project_identifier', 'N/A'))
        msg.attach(MIMEText(body, 'plain'))
        with open(pdf_path, 'rb') as f: part = MIMEApplication(f.read(), Name=os.path.basename(pdf_path))
        part['Content-Disposition'] = f'attachment; filename="{os.path.basename(pdf_path)}"'
        msg.attach(part)
        try:
            server = smtplib.SMTP(current_app_config["smtp_server"], current_app_config.get("smtp_port", 587))
            if current_app_config.get("smtp_port", 587) == 587: server.starttls()
            server.login(current_app_config["smtp_user"], current_app_config["smtp_password"])
            server.send_message(msg); server.quit()
            QMessageBox.information(self, self.tr("Email envoyé"), self.tr("Le document a été envoyé avec succès."))
        except Exception as e: QMessageBox.critical(self, self.tr("Erreur d'envoi"), self.tr("Erreur lors de l'envoi de l'email:
{0}").format(str(e)))

class ClientWidget(QWidget):
    def __init__(self, client_info, parent_config, main_window_ref=None): # parent_config is CONFIG from main_window
        super().__init__(main_window_ref) # main_window_ref is the DocumentManager instance
        self.client_info = client_info
        self.config = parent_config # Use the passed config
        self.main_window = main_window_ref # Store reference to main window for callbacks if needed
        self.setup_ui()

    manage_project_requested = pyqtSignal(str) # Signal takes client_id

    def setup_ui(self):
        layout = QVBoxLayout(self); layout.setContentsMargins(15, 15, 15, 15); layout.setSpacing(15)
        header = QLabel(f"<h2>{self.client_info['client_name']}</h2>"); header.setStyleSheet("color: #2c3e50;")
        layout.addWidget(header)
        action_layout = QHBoxLayout()
        self.create_docs_btn = QPushButton(self.tr("Créer Documents")); self.create_docs_btn.setIcon(QIcon.fromTheme("document-new")); self.create_docs_btn.setStyleSheet("background-color: #3498db; color: white;"); self.create_docs_btn.clicked.connect(self.open_create_docs_dialog_for_client) # Renamed
        action_layout.addWidget(self.create_docs_btn)
        self.compile_pdf_btn = QPushButton(self.tr("Compiler PDF")); self.compile_pdf_btn.setIcon(QIcon.fromTheme("document-export")); self.compile_pdf_btn.setStyleSheet("background-color: #27ae60; color: white;"); self.compile_pdf_btn.clicked.connect(self.open_compile_pdf_dialog_for_client) # Renamed
        action_layout.addWidget(self.compile_pdf_btn)

        self.manage_project_btn = QPushButton(self.tr("Manage Project Details"))
        self.manage_project_btn.setIcon(QIcon.fromTheme("preferences-system"))
        self.manage_project_btn.setStyleSheet("background-color: #545b62; color: white;")
        self.manage_project_btn.clicked.connect(self._emit_manage_project_request)
        action_layout.addWidget(self.manage_project_btn)
        action_layout.addStretch() # Ensure buttons align left if there's space

        layout.addLayout(action_layout)
        status_layout = QHBoxLayout(); status_label = QLabel(self.tr("Statut:")); status_layout.addWidget(status_label)
        self.status_combo = QComboBox()
        self.load_client_statuses_for_combo() # Renamed
        self.status_combo.setCurrentText(self.client_info.get("status", self.tr("En cours")))
        self.status_combo.currentTextChanged.connect(self.update_client_status_in_db) # Renamed
        self.status_combo.setStyleSheet(STYLE_GENERIC_INPUT)
        status_layout.addWidget(self.status_combo); layout.addLayout(status_layout)

        details_layout = QFormLayout(); details_layout.setLabelAlignment(Qt.AlignRight)
        details_data = [
            (self.tr("Référence Client/Commande:"), self.client_info.get("project_identifier", self.tr("N/A"))),
            (self.tr("Pays:"), self.client_info.get("country", self.tr("N/A"))),
            (self.tr("Ville:"), self.client_info.get("city", self.tr("N/A"))),
            (self.tr("Besoin Principal:"), self.client_info.get("need", self.tr("N/A"))),
            (self.tr("Prix Final:"), f"{self.client_info.get('price', 0)} €"),
            (self.tr("Date Création:"), self.client_info.get("creation_date", self.tr("N/A"))),
            (self.tr("Chemin Dossier:"), f"<a href='file:///{self.client_info['base_folder_path']}'>{self.client_info['base_folder_path']}</a>")
        ]
        for label_text, value_text in details_data:
            label_widget = QLabel(label_text); value_widget = QLabel(value_text)
            value_widget.setOpenExternalLinks(True); value_widget.setTextInteractionFlags(Qt.TextBrowserInteraction)
            details_layout.addRow(label_widget, value_widget)
        layout.addLayout(details_layout)

        notes_group = QGroupBox(self.tr("Notes")); notes_layout = QVBoxLayout(notes_group)
        self.notes_edit = QTextEdit(self.client_info.get("notes", ""))
        self.notes_edit.setStyleSheet(STYLE_GENERIC_INPUT)
        self.notes_edit.setPlaceholderText(self.tr("Ajoutez des notes sur ce client..."))
        self.notes_edit.textChanged.connect(self.save_client_notes_to_db) # Renamed
        notes_layout.addWidget(self.notes_edit); layout.addWidget(notes_group)

        self.tab_widget = QTabWidget()
        docs_tab = QWidget(); docs_layout = QVBoxLayout(docs_tab)
        self.doc_table = QTableWidget(); self.doc_table.setColumnCount(5)
        self.doc_table.setHorizontalHeaderLabels([self.tr("Nom"), self.tr("Type"), self.tr("Langue"), self.tr("Date"), self.tr("Actions")])
        self.doc_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.doc_table.setSelectionBehavior(QAbstractItemView.SelectRows); self.doc_table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        docs_layout.addWidget(self.doc_table)
        doc_btn_layout = QHBoxLayout()
        refresh_btn = QPushButton(self.tr("Actualiser")); refresh_btn.setIcon(QIcon.fromTheme("view-refresh")); refresh_btn.clicked.connect(self.populate_doc_table_for_client) # Renamed
        refresh_btn.setStyleSheet(STYLE_NEUTRAL_BUTTON)
        doc_btn_layout.addWidget(refresh_btn)
        # Removed Open and Delete buttons here, actions are in cell widgets
        docs_layout.addLayout(doc_btn_layout); self.tab_widget.addTab(docs_tab, self.tr("Documents"))

        contacts_tab = QWidget(); contacts_layout = QVBoxLayout(contacts_tab)
        self.contacts_list = QListWidget(); self.contacts_list.setAlternatingRowColors(True); self.contacts_list.itemDoubleClicked.connect(self.edit_client_contact) # Renamed
        contacts_layout.addWidget(self.contacts_list)
        contacts_btn_layout = QHBoxLayout()
        self.add_contact_btn = QPushButton(self.tr("Ajouter Contact")); self.add_contact_btn.setIcon(QIcon.fromTheme("contact-new")); self.add_contact_btn.clicked.connect(self.add_new_client_contact) # Renamed
        self.add_contact_btn.setStyleSheet(STYLE_PRIMARY_BUTTON)
        contacts_btn_layout.addWidget(self.add_contact_btn)
        self.edit_contact_btn = QPushButton(self.tr("Modifier Contact")); self.edit_contact_btn.setIcon(QIcon.fromTheme("document-edit")); self.edit_contact_btn.clicked.connect(self.edit_client_contact) # Renamed
        self.edit_contact_btn.setStyleSheet(STYLE_SECONDARY_BUTTON)
        contacts_btn_layout.addWidget(self.edit_contact_btn)
        self.remove_contact_btn = QPushButton(self.tr("Supprimer Contact")); self.remove_contact_btn.setIcon(QIcon.fromTheme("edit-delete")); self.remove_contact_btn.clicked.connect(self.remove_client_contact_link) # Renamed
        self.remove_contact_btn.setStyleSheet(STYLE_DANGER_BUTTON)
        contacts_btn_layout.addWidget(self.remove_contact_btn)
        contacts_layout.addLayout(contacts_btn_layout); self.tab_widget.addTab(contacts_tab, self.tr("Contacts"))

        products_tab = QWidget(); products_layout = QVBoxLayout(products_tab)
        self.products_table = QTableWidget(); self.products_table.setColumnCount(6)
        self.products_table.setHorizontalHeaderLabels([self.tr("ID"), self.tr("Nom Produit"), self.tr("Description"), self.tr("Qté"), self.tr("Prix Unitaire"), self.tr("Prix Total")])
        self.products_table.setEditTriggers(QAbstractItemView.NoEditTriggers); self.products_table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.products_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Stretch) # Name column stretch
        self.products_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.Stretch) # Description column stretch
        self.products_table.hideColumn(0)
        products_layout.addWidget(self.products_table)
        products_btn_layout = QHBoxLayout()
        self.add_product_btn = QPushButton(self.tr("Ajouter Produit")); self.add_product_btn.setIcon(QIcon.fromTheme("list-add")); self.add_product_btn.clicked.connect(self.add_product_to_client) # Renamed
        self.add_product_btn.setStyleSheet(STYLE_PRIMARY_BUTTON)
        products_btn_layout.addWidget(self.add_product_btn)
        self.edit_product_btn = QPushButton(self.tr("Modifier Produit")); self.edit_product_btn.setIcon(QIcon.fromTheme("document-edit")); self.edit_product_btn.clicked.connect(self.edit_linked_client_product) # Renamed
        self.edit_product_btn.setStyleSheet(STYLE_SECONDARY_BUTTON)
        products_btn_layout.addWidget(self.edit_product_btn)
        self.remove_product_btn = QPushButton(self.tr("Supprimer Produit")); self.remove_product_btn.setIcon(QIcon.fromTheme("edit-delete")); self.remove_product_btn.clicked.connect(self.remove_linked_client_product) # Renamed
        self.remove_product_btn.setStyleSheet(STYLE_DANGER_BUTTON)
        products_btn_layout.addWidget(self.remove_product_btn)
        products_layout.addLayout(products_btn_layout); self.tab_widget.addTab(products_tab, self.tr("Produits"))

        layout.addWidget(self.tab_widget)
        self.populate_doc_table_for_client() # Renamed
        self.load_contacts_for_client() # Renamed
        self.load_products_for_client() # Renamed

        # PDF Generation Section
        pdf_generation_group = QGroupBox(self.tr("Générer Documents PDF"))
        pdf_gen_layout = QVBoxLayout()

        self.btn_generate_proforma = QPushButton(self.tr("Générer Facture Proforma PDF"))
        self.btn_generate_proforma.setIcon(QIcon.fromTheme("document-export")) # Example icon
        self.btn_generate_proforma.clicked.connect(lambda: self.generate_document_pdf('proforma'))
        pdf_gen_layout.addWidget(self.btn_generate_proforma)

        self.btn_generate_packing_list = QPushButton(self.tr("Générer Packing List PDF"))
        self.btn_generate_packing_list.setIcon(QIcon.fromTheme("document-export"))
        self.btn_generate_packing_list.clicked.connect(lambda: self.generate_document_pdf('packing_list'))
        pdf_gen_layout.addWidget(self.btn_generate_packing_list)

        self.btn_generate_sales_contract = QPushButton(self.tr("Générer Contrat de Vente PDF"))
        self.btn_generate_sales_contract.setIcon(QIcon.fromTheme("document-export"))
        self.btn_generate_sales_contract.clicked.connect(lambda: self.generate_document_pdf('sales_contract'))
        pdf_gen_layout.addWidget(self.btn_generate_sales_contract)

        self.btn_generate_warranty = QPushButton(self.tr("Générer Document de Garantie PDF"))
        self.btn_generate_warranty.setIcon(QIcon.fromTheme("document-export"))
        self.btn_generate_warranty.clicked.connect(lambda: self.generate_document_pdf('warranty'))
        pdf_gen_layout.addWidget(self.btn_generate_warranty)

        self.btn_generate_cover_page = QPushButton(self.tr("Générer Page de Garde PDF"))
        self.btn_generate_cover_page.setIcon(QIcon.fromTheme("document-export"))
        self.btn_generate_cover_page.clicked.connect(lambda: self.generate_document_pdf('cover_page'))
        pdf_gen_layout.addWidget(self.btn_generate_cover_page)

        pdf_generation_group.setLayout(pdf_gen_layout)
        layout.addWidget(pdf_generation_group)


    def load_client_statuses_for_combo(self): # Renamed
        self.status_combo.clear()
        try:
            # Assuming db_manager.get_all_status_settings filters by type 'Client' or similar
            client_statuses = db_manager.get_all_status_settings(status_type='Client') # Assumed parameter
            if client_statuses is None: client_statuses = []
            for status_dict in client_statuses:
                self.status_combo.addItem(status_dict['status_name'], status_dict.get('status_id'))
        except Exception as e: # Catch generic db_manager error
            QMessageBox.warning(self, self.tr("Erreur DB"), self.tr("Erreur de chargement des statuts:
{0}").format(str(e)))

    def update_client_status_in_db(self, status_text): # Renamed
        status_id = self.status_combo.currentData()
        if status_id is None: return # Should not happen if combo is populated correctly

        try:
            updated = db_manager.update_client(self.client_info["client_id"], {'status_id': status_id})
            if updated:
                self.client_info["status"] = status_text # Update local cache for display
                self.client_info["status_id"] = status_id
                # Refresh client list in main window if main_window reference is available
                if self.main_window and hasattr(self.main_window, 'filter_client_list_display'):
                    self.main_window.filter_client_list_display()
                if self.main_window and hasattr(self.main_window, 'stats_widget'):
                    self.main_window.stats_widget.update_stats()
            else:
                 QMessageBox.warning(self, self.tr("Erreur DB"), self.tr("Échec de la mise à jour du statut."))
        except Exception as e:
            QMessageBox.warning(self, self.tr("Erreur DB"), self.tr("Erreur de mise à jour du statut:
{0}").format(str(e)))

    def save_client_notes_to_db(self): # Renamed
        notes = self.notes_edit.toPlainText()
        try:
            updated = db_manager.update_client(self.client_info["client_id"], {'notes': notes})
            if updated: self.client_info["notes"] = notes
            # else: QMessageBox.warning(self, self.tr("Erreur DB"), self.tr("Échec sauvegarde des notes.")) # Optional feedback
        except Exception as e:
            QMessageBox.warning(self, self.tr("Erreur DB"), self.tr("Erreur de sauvegarde des notes:
{0}").format(str(e)))

    def populate_doc_table_for_client(self): # Renamed
        self.doc_table.setRowCount(0)
        client_path = self.client_info["base_folder_path"]
        if not os.path.exists(client_path): return
        row = 0
        for lang_code_folder in self.client_info.get("selected_languages", ["fr"]):
            lang_dir = os.path.join(client_path, lang_code_folder)
            if not os.path.exists(lang_dir): continue
            for file_name in os.listdir(lang_dir):
                if file_name.lower().endswith(('.xlsx', '.pdf', '.docx', '.html')):
                    file_path = os.path.join(lang_dir, file_name)
                    name_item = QTableWidgetItem(file_name); name_item.setData(Qt.UserRole, file_path)
                    file_type_str = os.path.splitext(file_name)[1][1:].upper() # XLSX, PDF etc.
                    mod_time = datetime.fromtimestamp(os.path.getmtime(file_path)).strftime('%Y-%m-%d %H:%M')
                    self.doc_table.insertRow(row)
                    self.doc_table.setItem(row, 0, name_item)
                    self.doc_table.setItem(row, 1, QTableWidgetItem(file_type_str))
                    self.doc_table.setItem(row, 2, QTableWidgetItem(lang_code_folder))
                    self.doc_table.setItem(row, 3, QTableWidgetItem(mod_time))
                    action_widget = QWidget(); action_layout = QHBoxLayout(action_widget); action_layout.setContentsMargins(0,0,0,0); action_layout.setSpacing(5)
                    open_btn_i = QPushButton("📂"); open_btn_i.setToolTip(self.tr("Ouvrir")); open_btn_i.setFixedSize(28,28); open_btn_i.setStyleSheet("background-color: transparent; border: none; font-size: 16px;"); open_btn_i.clicked.connect(lambda _, p=file_path: self.open_document_for_client(p)); action_layout.addWidget(open_btn_i) # Renamed
                    if file_name.lower().endswith('.xlsx') or file_name.lower().endswith('.html'):
                        edit_btn_i = QPushButton("✏️"); edit_btn_i.setToolTip(self.tr("Éditer")); edit_btn_i.setFixedSize(28,28); edit_btn_i.setStyleSheet("background-color: transparent; border: none; font-size: 16px;"); edit_btn_i.clicked.connect(lambda _, p=file_path: self.open_document_for_client(p, edit_mode=True)); action_layout.addWidget(edit_btn_i) # Added edit_mode
                    delete_btn_i = QPushButton("🗑️"); delete_btn_i.setToolTip(self.tr("Supprimer")); delete_btn_i.setFixedSize(28,28); delete_btn_i.setStyleSheet("background-color: transparent; border: none; font-size: 16px;"); delete_btn_i.clicked.connect(lambda _, p=file_path: self.delete_document_for_client(p)); action_layout.addWidget(delete_btn_i) # Renamed
                    self.doc_table.setCellWidget(row, 4, action_widget)
                    row += 1
        self.doc_table.resizeColumnsToContents()


    def open_create_docs_dialog_for_client(self): # Renamed
        dialog = CreateDocumentDialog(self.client_info, self) # Uses global CONFIG
        if dialog.exec_() == QDialog.Accepted: self.populate_doc_table_for_client()

    def open_compile_pdf_dialog_for_client(self): # Renamed
        dialog = CompilePdfDialog(self.client_info, self)
        dialog.exec_() # Result handling (refresh) is done by CompilePdfDialog itself via accept()

    def open_document_for_client(self, file_path, edit_mode=False): # Renamed, added edit_mode
        if os.path.exists(file_path):
            try:
                editor_client_data = { # Data for populating templates/editors
                    "Nom du client": self.client_info.get("client_name", ""), "Besoin": self.client_info.get("need", ""),
                    "price": self.client_info.get("price", 0), "project_identifier": self.client_info.get("project_identifier", ""),
                    "company_name": self.client_info.get("company_name", ""), "country": self.client_info.get("country", ""),
                    "city": self.client_info.get("city", ""),
                    "client_id": self.client_info.get("client_id"), # Added client_id
                }
                if file_path.lower().endswith('.xlsx') and edit_mode:
                    editor = ExcelEditor(file_path, client_data=editor_client_data, parent=self)
                    if editor.exec_() == QDialog.Accepted: self.populate_doc_table_for_client()
                elif file_path.lower().endswith('.html') and edit_mode:
                    html_editor_dialog = HtmlEditor(file_path, client_data=editor_client_data, parent=self)
                    if html_editor_dialog.exec_() == QDialog.Accepted: self.populate_doc_table_for_client()
                else: # Default open for PDF, DOCX, or non-edit mode for XLSX/HTML
                    QDesktopServices.openUrl(QUrl.fromLocalFile(file_path))
            except Exception as e: QMessageBox.warning(self, self.tr("Erreur Ouverture Fichier"), self.tr("Impossible d'ouvrir/éditer le fichier:
{0}").format(str(e)))
        else:
            QMessageBox.warning(self, self.tr("Fichier Introuvable"), self.tr("Le fichier n'existe plus.")); self.populate_doc_table_for_client()

    def delete_document_for_client(self, file_path): # Renamed
        if not os.path.exists(file_path): return
        reply = QMessageBox.question(self, self.tr("Confirmer la suppression"), self.tr("Supprimer le fichier {0} ?").format(os.path.basename(file_path)), QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
        if reply == QMessageBox.Yes:
            try:
                os.remove(file_path); self.populate_doc_table_for_client()
                QMessageBox.information(self, self.tr("Fichier supprimé"), self.tr("Le fichier a été supprimé.")) # Simplified
            except Exception as e: QMessageBox.warning(self, self.tr("Erreur"), self.tr("Impossible de supprimer le fichier:
{0}").format(str(e)))

    def load_contacts_for_client(self): # Renamed
        self.contacts_list.clear()
        client_uuid = self.client_info.get("client_id");
        if not client_uuid: return
        try:
            contacts = db_manager.get_contacts_for_client(client_uuid)
            if contacts is None: contacts = []
            for contact in contacts:
                contact_text = f"{contact.get('name', 'N/A')}"
                if contact.get('email'): contact_text += f" <{contact.get('email')}>" # Added email
                if contact.get('phone'): contact_text += f" ({contact.get('phone')})"
                if contact.get('is_primary_for_client'): contact_text += f" [{self.tr('Principal')}]"
                item = QListWidgetItem(contact_text)
                item.setData(Qt.UserRole, {'contact_id': contact.get('contact_id'), 'client_contact_id': contact.get('client_contact_id')})
                self.contacts_list.addItem(item)
        except Exception as e: QMessageBox.warning(self, self.tr("Erreur DB"), self.tr("Erreur chargement contacts:
{0}").format(str(e)))

    def add_new_client_contact(self): # Renamed
        client_uuid = self.client_info.get("client_id");
        if not client_uuid: return
        dialog = ContactDialog(parent=self) # client_id not passed to dialog, it's for linking
        if dialog.exec_() == QDialog.Accepted:
            contact_form_data = dialog.get_data()
            try:
                existing_contact = db_manager.get_contact_by_email(contact_form_data['email'])
                contact_id_to_link = None
                if existing_contact:
                    contact_id_to_link = existing_contact['contact_id']
                    update_data = {k: v for k, v in contact_form_data.items() if k in ['name', 'phone', 'position'] and v != existing_contact.get(k)}
                    if update_data : db_manager.update_contact(contact_id_to_link, update_data)
                else:
                    new_contact_id = db_manager.add_contact({'name': contact_form_data['name'], 'email': contact_form_data['email'], 'phone': contact_form_data['phone'], 'position': contact_form_data['position']})
                    if new_contact_id: contact_id_to_link = new_contact_id
                    else: QMessageBox.critical(self, self.tr("Erreur DB"), self.tr("Impossible de créer contact global.")); return
                if contact_id_to_link:
                    if contact_form_data['is_primary']: # Handle unsetting other primaries
                        client_contacts = db_manager.get_contacts_for_client(client_uuid)
                        if client_contacts:
                            for cc in client_contacts:
                                if cc['is_primary_for_client']: db_manager.update_client_contact_link(cc['client_contact_id'], {'is_primary_for_client': False})
                    link_id = db_manager.link_contact_to_client(client_uuid, contact_id_to_link, is_primary=contact_form_data['is_primary'])
                    if link_id: self.load_contacts_for_client()
                    else: QMessageBox.critical(self, self.tr("Erreur DB"), self.tr("Impossible de lier contact au client (lien existe peut-être).")) # Simplified
            except Exception as e: QMessageBox.critical(self, self.tr("Erreur DB"), self.tr("Erreur d'ajout du contact:
{0}").format(str(e)))

    def edit_client_contact(self): # Renamed
        item = self.contacts_list.currentItem();
        if not item: return
        item_data = item.data(Qt.UserRole) # Contains contact_id and client_contact_id
        contact_id = item_data.get('contact_id')
        client_contact_id = item_data.get('client_contact_id')
        client_uuid = self.client_info.get("client_id")
        if not contact_id or not client_uuid: return
        try:
            contact_details = db_manager.get_contact_by_id(contact_id)
            if contact_details:
                current_link_info = db_manager.get_client_contact_link_details(client_contact_id) # Needs this function
                is_primary_for_this_client = current_link_info.get('is_primary_for_client', False) if current_link_info else False
                dialog_data = {"name": contact_details.get('name'), "email": contact_details.get('email'), "phone": contact_details.get('phone'), "position": contact_details.get('position'), "is_primary": is_primary_for_this_client}
                dialog = ContactDialog(contact_data=dialog_data, parent=self) # Pass contact_data for editing
                if dialog.exec_() == QDialog.Accepted:
                    new_form_data = dialog.get_data()
                    db_manager.update_contact(contact_id, {'name': new_form_data['name'], 'email': new_form_data['email'], 'phone': new_form_data['phone'], 'position': new_form_data['position']})
                    if new_form_data['is_primary'] != is_primary_for_this_client: # Primary status changed
                        if new_form_data['is_primary']: # Became primary
                            client_contacts = db_manager.get_contacts_for_client(client_uuid)
                            if client_contacts:
                                for cc in client_contacts: # Unset other primaries
                                    if cc['contact_id'] != contact_id and cc['is_primary_for_client']:
                                        db_manager.update_client_contact_link(cc['client_contact_id'], {'is_primary_for_client': False})
                        db_manager.update_client_contact_link(client_contact_id, {'is_primary_for_client': new_form_data['is_primary']})
                    self.load_contacts_for_client()
        except Exception as e: QMessageBox.critical(self, self.tr("Erreur DB"), self.tr("Erreur modification contact:
{0}").format(str(e)))

    def remove_client_contact_link(self): # Renamed
        item = self.contacts_list.currentItem();
        if not item: return
        item_data = item.data(Qt.UserRole); contact_id = item_data.get('contact_id')
        client_uuid = self.client_info.get("client_id")
        if not contact_id or not client_uuid: return
        contact_name = db_manager.get_contact_by_id(contact_id)['name'] if db_manager.get_contact_by_id(contact_id) else "Inconnu"
        reply = QMessageBox.question(self, self.tr("Confirmer Suppression Lien"), self.tr("Supprimer lien vers {0} pour ce client ?
(Contact global non supprimé).").format(contact_name), QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
        if reply == QMessageBox.Yes:
            try:
                unlinked = db_manager.unlink_contact_from_client(client_uuid, contact_id)
                if unlinked: self.load_contacts_for_client()
                else: QMessageBox.critical(self, self.tr("Erreur DB"), self.tr("Erreur suppression lien contact-client."))
            except Exception as e: QMessageBox.critical(self, self.tr("Erreur DB"), self.tr("Erreur suppression lien contact:
{0}").format(str(e)))

    def add_product_to_client(self): # Renamed
        client_uuid = self.client_info.get("client_id");
        if not client_uuid: return
        dialog = ProductDialog(client_uuid, parent=self) # client_uuid passed for context
        if dialog.exec_() == QDialog.Accepted:
            form_data = dialog.get_data() # name, desc, qty, unit_price_for_dialog
            try:
                global_product = db_manager.get_product_by_name(form_data['product_name'])
                global_product_id = None
                if global_product:
                    global_product_id = global_product['product_id']
                    # Optionally update global product if form_data['product_description'] differs
                else:
                    new_global_product_id = db_manager.add_product({'product_name': form_data['product_name'], 'description': form_data['product_description'], 'base_unit_price': form_data['unit_price_for_dialog']})
                    if new_global_product_id: global_product_id = new_global_product_id
                    else: QMessageBox.critical(self, self.tr("Erreur DB"), self.tr("Impossible de créer produit global.")); return
                if global_product_id:
                    # Determine if unit_price_for_dialog is an override
                    base_price = global_product['base_unit_price'] if global_product else form_data['unit_price_for_dialog']
                    price_override = form_data['unit_price_for_dialog'] if form_data['unit_price_for_dialog'] != base_price else None

                    link_data = {'client_id': client_uuid, 'project_id': None, 'product_id': global_product_id, 'quantity': form_data['quantity'], 'unit_price_override': price_override}
                    cpp_id = db_manager.add_product_to_client_or_project(link_data)
                    if cpp_id: self.load_products_for_client()
                    else: QMessageBox.critical(self, self.tr("Erreur DB"), self.tr("Impossible de lier produit au client."))
            except Exception as e: QMessageBox.critical(self, self.tr("Erreur DB"), self.tr("Erreur d'ajout du produit:
{0}").format(str(e)))

    def edit_linked_client_product(self): # Renamed
        selected_row = self.products_table.currentRow();
        if selected_row < 0: return
        cpp_id_item = self.products_table.item(selected_row, 0) # Hidden ID column
        if not cpp_id_item: return
        client_project_product_id = cpp_id_item.data(Qt.UserRole)
        client_uuid = self.client_info.get("client_id")
        try:
            linked_product_details = db_manager.get_client_project_product_by_id(client_project_product_id) # Needs this function
            if linked_product_details:
                # Data for ProductDialog: product_name, product_description, quantity, unit_price_override, base_unit_price
                dialog = ProductDialog(client_uuid, product_data=linked_product_details, parent=self)
                if dialog.exec_() == QDialog.Accepted:
                    form_data = dialog.get_data() # name, desc, qty, unit_price_for_dialog
                    # Update global product part
                    db_manager.update_product(linked_product_details['product_id'], {'product_name': form_data['product_name'], 'description': form_data['product_description']})
                    # Determine if unit_price_for_dialog is an override for the link
                    base_price = db_manager.get_product_by_id(linked_product_details['product_id'])['base_unit_price']
                    price_override = form_data['unit_price_for_dialog'] if form_data['unit_price_for_dialog'] != base_price else None
                    update_link_data = {'quantity': form_data['quantity'], 'unit_price_override': price_override}
                    if db_manager.update_client_project_product(client_project_product_id, update_link_data): self.load_products_for_client()
                    else: QMessageBox.critical(self, self.tr("Erreur DB"), self.tr("Échec mise à jour produit lié."))
            else: QMessageBox.warning(self, self.tr("Erreur"), self.tr("Détails produit lié introuvables."))
        except Exception as e: QMessageBox.critical(self, self.tr("Erreur DB"), self.tr("Erreur modification produit:
{0}").format(str(e)))

    def remove_linked_client_product(self): # Renamed
        selected_row = self.products_table.currentRow();
        if selected_row < 0: return
        cpp_id_item = self.products_table.item(selected_row, 0);
        if not cpp_id_item: return
        client_project_product_id = cpp_id_item.data(Qt.UserRole)
        product_name = self.products_table.item(selected_row, 1).text() if self.products_table.item(selected_row, 1) else self.tr("Inconnu")
        reply = QMessageBox.question(self, self.tr("Confirmer Suppression"), self.tr("Supprimer produit '{0}' de ce client?").format(product_name), QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
        if reply == QMessageBox.Yes:
            try:
                if db_manager.remove_product_from_client_or_project(client_project_product_id): self.load_products_for_client()
                else: QMessageBox.critical(self, self.tr("Erreur DB"), self.tr("Erreur suppression produit lié."))
            except Exception as e: QMessageBox.critical(self, self.tr("Erreur DB"), self.tr("Erreur suppression produit lié:
{0}").format(str(e)))

    def load_products_for_client(self): # Renamed
        self.products_table.setRowCount(0)
        client_uuid = self.client_info.get("client_id");
        if not client_uuid: return
        try:
            linked_products = db_manager.get_products_for_client_or_project(client_uuid, project_id=None)
            if linked_products is None: linked_products = []
            for row_idx, prod_link_data in enumerate(linked_products):
                self.products_table.insertRow(row_idx)
                id_item = QTableWidgetItem(str(prod_link_data.get('client_project_product_id'))); id_item.setData(Qt.UserRole, prod_link_data.get('client_project_product_id')); self.products_table.setItem(row_idx, 0, id_item)
                self.products_table.setItem(row_idx, 1, QTableWidgetItem(prod_link_data.get('product_name', 'N/A')))
                self.products_table.setItem(row_idx, 2, QTableWidgetItem(prod_link_data.get('product_description', '')))
                qty_item = QTableWidgetItem(str(prod_link_data.get('quantity', 0))); qty_item.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter); self.products_table.setItem(row_idx, 3, qty_item)
                effective_unit_price = prod_link_data.get('unit_price_override', prod_link_data.get('base_unit_price', 0.0))
                unit_price_item = QTableWidgetItem(f"€ {effective_unit_price:.2f}"); unit_price_item.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter); self.products_table.setItem(row_idx, 4, unit_price_item)
                total_price_item = QTableWidgetItem(f"€ {prod_link_data.get('total_price_calculated', 0.0):.2f}"); total_price_item.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter); self.products_table.setItem(row_idx, 5, total_price_item)
            self.products_table.resizeColumnsToContents()
        except Exception as e: QMessageBox.warning(self, self.tr("Erreur DB"), self.tr("Erreur chargement produits:
{0}").format(str(e)))

def populate_docx_template(docx_path, client_data):
    try:
        document = Document(docx_path)
        placeholders = {
            "{{CLIENT_NAME}}": client_data.get('client_name', ''), "{{PROJECT_ID}}": client_data.get('project_identifier', ''),
            "{{COMPANY_NAME}}": client_data.get('company_name', ''), "{{NEED}}": client_data.get('need', ''),
            "{{COUNTRY}}": client_data.get('country', ''), "{{CITY}}": client_data.get('city', ''),
            "{{PRICE}}": str(client_data.get('price', 0)), "{{DATE}}": datetime.now().strftime('%Y-%m-%d'),
            "{{STATUS}}": client_data.get('status', ''),
            "{{SELECTED_LANGUAGES}}": ", ".join(client_data.get('selected_languages', [])),
            "{{NOTES}}": client_data.get('notes', ''), "{{CREATION_DATE}}": client_data.get('creation_date', ''),
            "{{CATEGORY}}": client_data.get('category', ''), "{{PRIMARY_CONTACT_NAME}}": "",
        }
        for para in document.paragraphs:
            for key, value in placeholders.items():
                if key in para.text: para.text = para.text.replace(key, value)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for key, value in placeholders.items():
                            if key in para.text: para.text = para.text.replace(key, value)
        document.save(docx_path)
    except Exception as e:
        print(f"Error populating DOCX template {docx_path}: {e}")
        raise

    def generate_document_pdf(self, doc_type_key: str):
        template_map = {
            'proforma': {'filename': 'proforma_invoice_template.html', 'title': self.tr("Facture Proforma")},
            'packing_list': {'filename': 'packing_list_template.html', 'title': self.tr("Packing List")},
            'sales_contract': {'filename': 'sales_contract_template.html', 'title': self.tr("Contrat de Vente")},
            'warranty': {'filename': 'warranty_document_template.html', 'title': self.tr("Document de Garantie")},
            'cover_page': {'filename': 'cover_page_template.html', 'title': self.tr("Page de Garde")},
        }

        if doc_type_key not in template_map:
            QMessageBox.warning(self, self.tr("Erreur"), self.tr("Type de document inconnu."))
            return

        template_info = template_map[doc_type_key]
        template_filename = template_info['filename']
        doc_title_for_context = template_info['title']

        client_id = self.client_info.get('client_id')
        # Prefer 'project_id' (UUID) if available, else 'project_identifier' (human-readable)
        # db_manager.get_document_context_data expects project_id to be the UUID if provided.
        project_id = self.client_info.get('project_id') or self.client_info.get('project_identifier')


        if not client_id:
            QMessageBox.warning(self, self.tr("Erreur Client"), self.tr("ID Client non trouvé."))
            return

        default_company = db_manager.get_default_company()
        if not default_company:
            QMessageBox.warning(self, self.tr("Erreur Configuration"), self.tr("Aucune société par défaut configurée."))
            return
        company_id = default_company['company_id']

        selected_languages = self.client_info.get("selected_languages", ["fr"]) # Default to 'fr'
        lang_code = selected_languages[0] if selected_languages and isinstance(selected_languages, list) else "fr"


        template_path = os.path.join(self.config.get("templates_dir", "templates"), lang_code, template_filename)

        if not os.path.exists(template_path):
            QMessageBox.critical(self, self.tr("Erreur Fichier"), self.tr("Le fichier modèle HTML est introuvable: {0}").format(template_path))
            return

        try:
            with open(template_path, 'r', encoding='utf-8') as f:
                template_content = f.read()
        except Exception as e:
            QMessageBox.critical(self, self.tr("Erreur Lecture Modèle"), self.tr("Impossible de lire le fichier modèle HTML: {0}").format(str(e)))
            return

        additional_doc_context = {
            "document_title": doc_title_for_context,
            # Add any other specific context needed per doc_type_key
        }
        # Example of adding specific context based on document type
        current_time_str = datetime.now().strftime("%Y%m%d%H%M%S") # Added seconds for more uniqueness
        project_identifier_short = self.client_info.get('project_identifier', client_id[:5])
        today_date = datetime.now().strftime("%Y-%m-%d")

        # Default placeholders for complex/legal text
        default_claus_text = "[Specify Details Here]"
        default_terms_text = "[Specify Terms Here]"
        default_list_text = "[List Items Here]"

        if doc_type_key == 'proforma':
            additional_doc_context.update({
                "invoice_number": f"PRO-{project_identifier_short}-{current_time_str}",
                "invoice_date_issue": today_date,
                "invoice_date_due": (datetime.now() + timedelta(days=30)).strftime("%Y-%m-%d"),
                "payment_terms": "Paiement à 30 jours nets.",
                "proforma_validity_days": "30",
                "subtotal_formatted": "€ 0.00 (Calculé)", # Example, real calculation would be more complex
                "tax_description": "TVA",
                "tax_rate": "20",
                "tax_amount_formatted": "€ 0.00 (Calculé)",
                "grand_total_formatted": "€ 0.00 (Calculé)",
                "signature_name_and_title": "[Nom et Titre du Signataire]",
                # Bank details will primarily come from DB if structured, otherwise N/A or raw.
            })
        elif doc_type_key == 'packing_list':
             additional_doc_context.update({
                "packing_list_number": f"PL-{project_identifier_short}-{current_time_str}",
                "packing_list_date": today_date,
                "invoice_number": f"INV-{project_identifier_short}-{current_time_str}", # Link to a proforma/invoice
                "order_number": self.client_info.get('project_identifier', default_claus_text),
                "date_of_shipment": today_date,
                "port_of_loading": "[Port de Chargement]",
                "port_of_discharge": "[Port de Déchargement]",
                "vessel_flight_details": "[Détails Navire/Vol]",
                "shipping_marks": "[Marques d'Expédition]",
                "item_number_or_code": "[Code Article]",
                "item_quantity_packages": "1",
                "item_quantity_units_per_package": "1",
                "item_total_units": "1",
                "item_net_weight_per_unit_formatted": "0.00 kg",
                "item_gross_weight_per_unit_formatted": "0.00 kg",
                "item_total_net_weight_formatted": "0.00 kg",
                "item_total_gross_weight_formatted": "0.00 kg",
                "item_package_dimensions": "N/A",
                "item_package_volume_formatted": "0.00 m³",
                "total_packages_count": "0",
                "total_net_weight_shipment_formatted": "0.00 kg",
                "total_gross_weight_shipment_formatted": "0.00 kg",
                "total_volume_shipment_formatted": "0.00 m³",
                "declaration_statement": "[Déclaration Standard]",
                "instructions": "[Instructions Spéciales]",
                "signature_date": today_date,
                "notify_party_name": "[Nom Partie à Notifier]",
                "notify_party_address": "[Adresse Partie à Notifier]",
             })
        elif doc_type_key == 'sales_contract':
            additional_doc_context.update({
                "contract_number": f"SC-{project_identifier_short}-{current_time_str}",
                "date_of_agreement": today_date,
                "preamble_text": "[Préambule du contrat à spécifier]",
                "subject_matter_general_description": self.client_info.get('need', default_claus_text),
                "total_quantity": default_claus_text, # Example: "As per Annex A"
                "quality_specifications_details": default_claus_text,
                "total_price_formatted": "€ 0.00 (Calculé)",
                "total_price_in_words": "[Montant Total en Toutes Lettres]",
                "price_per_unit_formatted": "€ 0.00",
                "payment_method_description": default_terms_text,
                "payment_schedule_description": default_terms_text,
                "delivery_incoterms": "EXW (Incoterms® 2020)",
                "delivery_place": "[Lieu de Livraison]",
                "estimated_delivery_date": (datetime.now() + timedelta(days=45)).strftime("%Y-%m-%d"),
                "partial_shipments_allowed_description": "Non autorisées",
                "shipping_documents_required": default_list_text,
                "inspection_of_goods_clause": default_claus_text,
                "warranty_period": "1 an à compter de la date de livraison",
                "warranty_conditions": default_claus_text,
                "force_majeure_clause": default_claus_text,
                "applicable_law_jurisdiction": "[Loi Applicable et Juridiction]",
                "dispute_resolution_method": "[Méthode de Résolution des Litiges]",
                "seller_signature_date": today_date,
                "buyer_signature_date": today_date,
                "annexes_list": [], # Expect a list of {'annex_title': '...'}
                "annexes_placeholder_text": "Aucune annexe pour le moment."
            })
        elif doc_type_key == 'warranty':
            additional_doc_context.update({
                "document_title": self.tr("Certificat de Garantie"), # Override default title
                "related_invoice_number": f"INV-{project_identifier_short}-{current_time_str}",
                "date_of_issue_of_warranty": today_date,
                "date_of_purchase": today_date, # Assume purchase on issue date for default
                # products_covered would be a list of dicts: {'product_name':'X', 'product_serial_number':'Y', ...}
                "products_covered": [{'product_name': '[Nom du Produit Garanti]', 'product_serial_number': '[Numéro de Série]', 'product_other_identifier': '[Autre ID]'}],
                "product_name_single": "[Nom du Produit Garanti]", # For single product warranty
                "product_serial_number_single": "[Numéro de Série]",
                "product_description_single": "[Description du Produit]",
                "warranty_period_description": "Un (1) an à compter de la date d'achat.",
                "warranty_start_date": today_date,
                "warranty_end_date": (datetime.now() + timedelta(days=365)).strftime("%Y-%m-%d"),
                "warranty_coverage_details": default_claus_text,
                "exclusions_from_warranty_details": default_claus_text,
                "claim_procedure_details": default_claus_text,
                "remedies_details": default_claus_text,
                "limitation_of_liability_details": default_claus_text,
                "governing_law_jurisdiction": "[Loi Applicable - Garantie]",
                "signature_date": today_date,
            })
        elif doc_type_key == 'cover_page':
            additional_doc_context.update({
                "document_subtitle": self.client_info.get('project_identifier', self.tr("Rapport de Projet")),
                "document_version": "1.0",
                "document_date": today_date, # Overrides context["doc"]["current_date"] if used for cover page
            })

        # Fetch the main context, merging with additional specifics
        context = db_manager.get_document_context_data(client_id, company_id, project_id, additional_context=additional_doc_context)

        try:
            populated_html = render_html_template(template_content, context)
        except Exception as e:
            QMessageBox.critical(self, self.tr("Erreur Rendu HTML"), self.tr("Erreur lors du rendu du modèle HTML: {0}").format(str(e)))
            return

        # Determine base_url for WeasyPrint.
        # It needs to allow access to resources like company_logos.
        # APP_ROOT_DIR should be the absolute path to the application's root.
        # Example: if logos are in APP_ROOT_DIR/company_logos/
        base_url_for_pdf = f"file://{APP_ROOT_DIR}/"
        # Ensure APP_ROOT_DIR ends with a slash if it's used as a base for relative paths in HTML
        # However, os.path.join in get_document_context_data for logo_path_absolute should create correct absolute file paths.
        # So, base_url might be more for CSS or other template-relative assets not covered by absolute paths.
        # If all image paths (like logos) are made absolute in the context, base_url is less critical for them.
        # For now, using APP_ROOT_DIR as base_url.

        try:
            pdf_bytes = convert_html_to_pdf(populated_html, base_url=base_url_for_pdf)
            if not pdf_bytes:
                QMessageBox.warning(self, self.tr("Erreur PDF"), self.tr("La génération du PDF a échoué (contenu vide)."))
                return
        except Exception as e:
            QMessageBox.critical(self, self.tr("Erreur Génération PDF"), self.tr("Erreur lors de la conversion HTML en PDF: {0}").format(str(e)))
            return

        default_pdf_name = f"{doc_type_key}_{self.client_info.get('client_name', 'client').replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.pdf"

        # Ensure client_info["base_folder_path"] exists.
        client_base_folder = self.client_info.get("base_folder_path")
        if not client_base_folder or not os.path.isdir(client_base_folder):
            # Fallback to documents location if client's base folder is not set or invalid
            client_base_folder = QStandardPaths.writableLocation(QStandardPaths.DocumentsLocation)
            os.makedirs(client_base_folder, exist_ok=True) # Ensure this fallback directory exists

        client_lang_folder = os.path.join(client_base_folder, lang_code)
        os.makedirs(client_lang_folder, exist_ok=True)

        suggested_path = os.path.join(client_lang_folder, default_pdf_name)

        save_path, _ = QFileDialog.getSaveFileName(
            self, self.tr("Enregistrer PDF sous"), suggested_path, self.tr("Fichiers PDF (*.pdf)")
        )

        if save_path:
            try:
                with open(save_path, 'wb') as f:
                    f.write(pdf_bytes)
                QMessageBox.information(self, self.tr("Succès"), self.tr("PDF généré et sauvegardé: {0}").format(save_path))
                QDesktopServices.openUrl(QUrl.fromLocalFile(save_path))

                # Refresh document list in ClientWidget if it's displayed
                if hasattr(self, 'populate_doc_table_for_client'):
                    self.populate_doc_table_for_client()
            except Exception as e:
                QMessageBox.critical(self, self.tr("Erreur Sauvegarde"), self.tr("Impossible de sauvegarder le PDF: {0}").format(str(e)))
