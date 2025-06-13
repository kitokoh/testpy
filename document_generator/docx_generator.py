import os
import logging
import shutil
from docx import Document
from docx.shared import Inches # For potential image handling later, not used yet
from docx.enum.text import WD_ALIGN_PARAGRAPH # For potential formatting later

# Assuming context_builder is in the same package
from .context_builder import get_document_context

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Placeholder format: {{KEY}}
# This matches the keys in the context['placeholders'] dictionary.

def replace_placeholders_in_run(run, placeholders: dict):
    """
    Replaces placeholders in a single run of text.
    A run is a contiguous block of text with the same formatting.
    Placeholders can span across multiple runs if not careful, but python-docx
    processes text in runs. This function handles multiple placeholders within one run.
    """
    text = run.text
    for key, value in placeholders.items():
        placeholder = f"{{{{{key}}}}}" # e.g. {{CLIENT_NAME}}
        if placeholder in text:
            # Ensure value is a string, convert if necessary (e.g. numbers, dates)
            str_value = str(value) if value is not None else ""
            text = text.replace(placeholder, str_value)
    # Only update run.text if it actually changed to avoid clearing formatting
    # on runs that had no placeholders.
    if run.text != text:
        run.text = text


def replace_placeholders_in_paragraph(paragraph, placeholders: dict):
    """
    Iterates through runs in a paragraph and replaces placeholders.
    """
    for run in paragraph.runs:
        replace_placeholders_in_run(run, placeholders)

def generate_docx_document(
    client_id: int,
    company_id: int,
    template_docx_path: str,
    output_docx_path: str,
    language_code: str,
    project_id: int = None,
    document_type: str = None, # Added to pass to get_document_context
    additional_context: dict = None,
    db_conn=None
) -> bool:
    """
    Generates a DOCX document by populating a template with context data.

    Args:
        client_id: ID of the client.
        company_id: ID of the seller company.
        template_docx_path: Absolute path to the source .docx template file.
        output_docx_path: Absolute path where the generated .docx file will be saved.
        language_code: Language code for translations.
        project_id: Optional ID of the project.
        document_type: Optional type of document (e.g. 'proforma_invoice') for context.
        additional_context: Optional dictionary for overrides or extra information.
        db_conn: Optional existing database connection.

    Returns:
        True on success, False on failure.
    """
    if not os.path.exists(template_docx_path):
        logger.error(f"DOCX template not found at: {template_docx_path}")
        return False

    if additional_context is None:
        additional_context = {}

    effective_document_type = document_type or additional_context.get('document_type')

    logger.info(f"Generating DOCX document: client_id={client_id}, company_id={company_id}, "
                f"template='{os.path.basename(template_docx_path)}', output='{os.path.basename(output_docx_path)}'")

    try:
        # 1. Fetch document context
        context = get_document_context(
            client_id=client_id,
            company_id=company_id,
            language_code=language_code,
            project_id=project_id,
            document_type=effective_document_type,
            additional_data=additional_context,
            db_conn=db_conn
        )
        if not context or "placeholders" not in context:
            logger.error("Failed to retrieve document context or placeholders missing.")
            return False

        placeholders = context.get("placeholders", {})
        # Ensure all placeholder values are strings, as python-docx expects strings for text runs.
        for key, value in placeholders.items():
            if value is None:
                placeholders[key] = "" # Replace None with empty string
            elif not isinstance(value, str):
                placeholders[key] = str(value)


        # 2. Copy template to output path to work on a copy
        output_dir = os.path.dirname(output_docx_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
            logger.info(f"Created output directory: {output_dir}")

        shutil.copy(template_docx_path, output_docx_path)

        # 3. Open the copied DOCX file
        doc = Document(output_docx_path)

        # 4. Replace placeholders in main body paragraphs
        logger.debug("Replacing placeholders in body paragraphs...")
        for para in doc.paragraphs:
            replace_placeholders_in_paragraph(para, placeholders)

        # 5. Replace placeholders in tables
        logger.debug("Replacing placeholders in tables...")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_placeholders_in_paragraph(para, placeholders)

        # 6. Replace placeholders in headers and footers
        # Headers
        logger.debug("Replacing placeholders in headers...")
        for section in doc.sections:
            # Header
            header = section.header
            if header:
                for para in header.paragraphs:
                    replace_placeholders_in_paragraph(para, placeholders)
                for table in header.tables: # Also check tables in header
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                replace_placeholders_in_paragraph(para, placeholders)
            # Footer
            footer = section.footer
            if footer:
                for para in footer.paragraphs:
                    replace_placeholders_in_paragraph(para, placeholders)
                for table in footer.tables: # Also check tables in footer
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                replace_placeholders_in_paragraph(para, placeholders)

            # Different First Page Header/Footer
            if section.different_first_page:
                first_page_header = section.first_page_header
                if first_page_header:
                    for para in first_page_header.paragraphs:
                         replace_placeholders_in_paragraph(para, placeholders)
                    for table in first_page_header.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    replace_placeholders_in_paragraph(para, placeholders)

                first_page_footer = section.first_page_footer
                if first_page_footer:
                    for para in first_page_footer.paragraphs:
                         replace_placeholders_in_paragraph(para, placeholders)
                    for table in first_page_footer.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    replace_placeholders_in_paragraph(para, placeholders)


        # Note: python-docx does not directly support text boxes easily.
        # Manipulation of shapes/text boxes usually requires deeper XML manipulation.

        # 7. Save the modified DOCX file
        doc.save(output_docx_path)
        logger.info(f"DOCX document generated and saved to: {output_docx_path}")
        return True

    except FileNotFoundError:
        logger.error(f"DOCX template file not found at: {template_docx_path} (should have been caught earlier)")
        return False
    except Exception as e:
        logger.error(f"Error generating DOCX document: {e}", exc_info=True)
        # Clean up potentially corrupted output file
        if os.path.exists(output_docx_path):
            try:
                os.remove(output_docx_path)
                logger.info(f"Cleaned up partially generated file: {output_docx_path}")
            except OSError as oe:
                logger.error(f"Error cleaning up file {output_docx_path}: {oe}")
        return False


if __name__ == '__main__':
    logger.info("Starting DOCX generator test run...")

    # This test requires:
    # 1. `context_builder.py` to be functional (or mocked).
    # 2. A sample .docx template file with placeholders like {{CLIENT_NAME}}.
    # 3. `python-docx` library installed.
    # 4. A running database instance or full mocks for context_builder.

    # Create a dummy template directory and output directory
    sample_dir = "temp_docx_test_files"
    os.makedirs(sample_dir, exist_ok=True)
    sample_template_path = os.path.join(sample_dir, "test_template.docx")
    sample_output_path = os.path.join(sample_dir, "test_output.docx")

    # Create a very basic DOCX template for testing
    try:
        doc = Document()
        doc.add_heading('Test Document for {{CLIENT_COMPANY_NAME}}', level=1)
        doc.add_paragraph('Hello, {{BUYER_CONTACT_NAME}}!')
        doc.add_paragraph('This document is prepared by {{SELLER_COMPANY_NAME}}.')
        doc.add_paragraph('Date: {{CURRENT_DATE}}')

        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Project ID"
        table.cell(0, 1).text = "{{PROJECT_IDENTIFIER}}"
        table.cell(1, 0).text = "Total Amount"
        table.cell(1, 1).text = "{{GRAND_TOTAL_AMOUNT}} {{CURRENCY_SYMBOL}}"

        # Add header and footer
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.text = "Header: {{SELLER_COMPANY_NAME}} - {{PROJECT_NAME}}"

        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.text = "Footer: Page X - {{DOC_TYPE_PLACEHOLDER}}" # DOC_TYPE_PLACEHOLDER as example

        doc.save(sample_template_path)
        logger.info(f"Sample DOCX template created at: {sample_template_path}")

    except Exception as e:
        logger.error(f"Failed to create sample DOCX template: {e}")
        # If template creation fails, the test cannot proceed meaningfully.
        # Consider exiting or skipping the test.

    # --- MOCKING (similar to html_generator, adapt as needed) ---
    # For this test, we'll rely on the actual context_builder if possible,
    # or a simplified mock if direct patching is done.
    # Assume get_document_context is either working with a test DB or mocked externally.

    # Example mock placeholders that `get_document_context` might return
    # This would typically be done by mocking `get_document_context` itself.
    # For this test, if `get_document_context` is not fully mocked,
    # it will try to connect to DB and might fail.
    # To ensure test runs, here's a conceptual override for this specific run:

    _original_get_document_context = get_document_context # keep a reference
    def _mock_get_document_context_for_docx(client_id, company_id, language_code, project_id=None, document_type=None, additional_data=None, db_conn=None):
        logger.info("[MOCK] _mock_get_document_context_for_docx called")
        # This should match the structure of `context_builder.get_document_context`
        return {
            "doc": {"current_date": "2023-10-26", "currency_symbol": "$"},
            "client": {"company_name": "Awesome Client LLC", "contact_name": "Ms. Doe"},
            "seller": {"name": "Reliable Seller Inc."},
            "project": {"name": "Super Project X", "identifier": "PROJ001"},
            "products": [], # Not directly used by placeholders but part of context
            "lang": {"code": language_code},
            "placeholders": {
                "CLIENT_COMPANY_NAME": "Awesome Client LLC",
                "BUYER_CONTACT_NAME": "Ms. Jane Doe",
                "SELLER_COMPANY_NAME": "Reliable Seller Inc.",
                "CURRENT_DATE": "2023-10-26",
                "PROJECT_IDENTIFIER": "PROJ001",
                "PROJECT_NAME": "Super Project X",
                "GRAND_TOTAL_AMOUNT": "1,234.56",
                "CURRENCY_SYMBOL": "$",
                "DOC_TYPE_PLACEHOLDER": "Test Proposal Document" # For footer
            },
            "additional": additional_data or {}
        }

    # Apply the mock for the test run
    # This is a direct assignment to the name used in this file.
    get_document_context = _mock_get_document_context_for_docx


    # --- Test generate_docx_document ---
    logger.info(f"\n--- Testing generate_docx_document ---")

    success = generate_docx_document(
        client_id=123,
        company_id=456,
        template_docx_path=sample_template_path,
        output_docx_path=sample_output_path,
        language_code='en',
        project_id=789,
        document_type='proposal', # example document type
        additional_context={"custom_field": "Custom Value for DOCX"}
    )

    if success:
        logger.info(f"generate_docx_document succeeded. Output at: {sample_output_path}")
        # You would typically open the DOCX file to manually verify content.
        # Automated verification would involve reading the DOCX and checking text,
        # which is more complex.
    else:
        logger.error("generate_docx_document failed.")

    # --- Cleanup ---
    # Restore original get_document_context
    get_document_context = _original_get_document_context

    # Optionally remove test files (uncomment to clean up)
    # if os.path.exists(sample_template_path):
    #     os.remove(sample_template_path)
    # if os.path.exists(sample_output_path):
    #     os.remove(sample_output_path)
    # if os.path.exists(sample_dir) and not os.listdir(sample_dir): # Remove dir if empty
    #     os.rmdir(sample_dir)

    logger.info("DOCX generator test run finished.")
    # Note: Requires `pip install python-docx`
    # The quality of placeholder replacement in headers/footers depends on
    # how they are structured in the template and `python-docx` capabilities.
    # Text boxes are generally not handled by this basic replacement logic.

```
